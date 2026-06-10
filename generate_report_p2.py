"""
Certichain Africa - FYP Report Generator Part 2
Covers: Chapter 4 Part 1 (Implementation), Chapter 4 Part 2 (Discussion),
        Chapter 5, Appendices, References
Then merges with Part 1 into final FYP_REPORT.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ── Load Part 1 doc ────────────────────────────────────────────────────────────
doc = Document('docs/FYP_REPORT_PART1.docx')

def add_body(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(12)
    run.font.bold   = bold
    run.font.italic = italic
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(text.upper())
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(14)
    run.font.bold  = True

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(12)
    run.font.bold  = True

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(12)
    run.font.bold   = True
    run.font.italic = True

def add_code(doc, filename, code):
    p = doc.add_paragraph()
    run = p.add_run(f"File: {filename}")
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.italic = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.left_indent  = Cm(1.0)
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(8)
    p2.paragraph_format.line_spacing = Pt(16)
    run2 = p2.add_run(code)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(9)

def add_fig(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(f'[{text}]')
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

def page_break(doc):
    doc.add_page_break()

def bullet(doc, text, indent=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(f"- {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ==============================================================================
# CHAPTER 4 PART 1: IMPLEMENTATION AND RESULTS
# ==============================================================================

add_heading1(doc, "CHAPTER 4 (PART 1)\nIMPLEMENTATION AND RESULTS")

add_heading2(doc, "4.1 System Analysis")

add_heading3(doc, "4.1.1 Problem Analysis")

add_body(doc, "The analysis phase began with a close examination of the problem that Certichain Africa is intended to solve. As established in the literature review, the core problem is that African educational institutions lack a reliable, affordable, and locally adapted system for issuing verifiable digital certificates. This problem was broken down into its component parts to identify exactly what the system would need to do in order to address it.")

add_body(doc, "The first component is certificate generation. Institutions need to be able to create certificates that look professional and that carry enough information to be meaningful to a recipient or a verifier. The certificate must include the recipient's name, the type of credential, the field of study or competency, the date of issuance, the issuing institution's name, and a unique identifier that can be used for verification. The certificate must be downloadable as a PDF file and must be visually distinctive enough that a verifier can reasonably tell it apart from a simple word-processed document.")

add_body(doc, "The second component is certificate anchoring. A certificate that exists only in the platform's own database is only as trustworthy as the platform operator. If the platform operator were to go offline, change ownership, or be compromised, the certificates stored in its database would lose their verifiability. Anchoring each certificate to a public blockchain creates an independent record that cannot be altered by the platform operator and that remains accessible even if the platform itself ceases to exist.")

add_body(doc, "The third component is certificate verification. Making certificates verifiable is only useful if the verification process is simple enough that any third party can perform it. A complex or technical verification process that requires the verifier to interact with blockchain tools would not be adopted in practice. The verification mechanism must be as simple as scanning a QR code with a smartphone or uploading a file to a web page.")

add_body(doc, "The fourth component is access control. The authority to issue certificates on behalf of an institution must be protected. If any user could log in and issue certificates, the entire system would be worthless. The authentication mechanism must be strong enough to prevent unauthorised access, including through stolen passwords, and must be usable by institution administrators who may not have extensive technical knowledge.")

add_body(doc, "The fifth component is payment. The platform must be financially sustainable. A freemium model with tiered subscription levels allows small institutions to use the platform at no cost while generating revenue from larger institutions. The payment mechanism must be one that is actually usable by institutions in Cameroon and the wider Central African region.")

add_heading3(doc, "4.1.2 Functional Requirements")

add_body(doc, "Based on the problem analysis, the following functional requirements were defined for Certichain Africa. Each requirement describes a capability that the system must provide.")

tbl = doc.add_table(rows=15, cols=3)
tbl.style = 'Table Grid'
for i, h in enumerate(["ID", "Requirement", "Priority"]):
    tbl.rows[0].cells[i].text = h
    for p in tbl.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

fr = [
    ("FR-01", "The system shall allow an institution to register a new account with a name, email, and password.", "High"),
    ("FR-02", "The system shall validate that the email address is unique and that the password meets the defined strength criteria.", "High"),
    ("FR-03", "The system shall authenticate institution users through email and password, followed by a one-time password delivered by email.", "High"),
    ("FR-04", "The system shall lock out an institution account after five consecutive incorrect OTP attempts.", "High"),
    ("FR-05", "The system shall allow an institution to reset its password through a secure token-based email link.", "Medium"),
    ("FR-06", "The system shall allow authenticated institutions to create a new certificate by specifying the type, recipient name, recipient email, domain, mention, and other relevant fields.", "High"),
    ("FR-07", "The system shall generate a professionally designed PDF certificate for each of the four supported types: diploma, certification, badge, and cyber diploma.", "High"),
    ("FR-08", "The system shall embed a QR code in each generated PDF that links to a public verification page for that certificate.", "High"),
    ("FR-09", "The system shall attempt to anchor each certificate to the Polygon blockchain, and fall back to SHA-256 local hashing if blockchain anchoring fails.", "High"),
    ("FR-10", "The system shall upload each certificate PDF to Pinata IPFS, and fall back to local file storage if IPFS upload fails.", "Medium"),
    ("FR-11", "The system shall allow authenticated institutions to view, download, and delete their certificates.", "High"),
    ("FR-12", "The system shall allow any user without an account to verify a certificate by uploading the PDF file or entering the blockchain hash.", "High"),
    ("FR-13", "The system shall allow institutions to initiate a MoMo payment to upgrade their subscription plan.", "Medium"),
    ("FR-14", "The system shall automatically upgrade an institution's plan when a MoMo payment is confirmed as successful.", "Medium"),
]

for i, (fid, req, pri) in enumerate(fr, start=1):
    tbl.rows[i].cells[0].text = fid
    tbl.rows[i].cells[1].text = req
    tbl.rows[i].cells[2].text = pri
    for j in range(3):
        for p in tbl.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

add_fig(doc, "Table 4.1: Functional Requirements of Certichain Africa")

add_heading3(doc, "4.1.3 Non-Functional Requirements")

tbl2 = doc.add_table(rows=9, cols=3)
tbl2.style = 'Table Grid'
for i, h in enumerate(["ID", "Requirement", "Category"]):
    tbl2.rows[0].cells[i].text = h
    for p in tbl2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

nfr = [
    ("NFR-01", "Passwords must be hashed using PBKDF2-SHA256 with a random salt before storage.", "Security"),
    ("NFR-02", "The OTP must expire within ten minutes of generation and must be stored as a SHA-256 hash in the session.", "Security"),
    ("NFR-03", "All routes that display or modify institution data must require a valid session.", "Security"),
    ("NFR-04", "The system must generate a certificate PDF in under three seconds on average server hardware.", "Performance"),
    ("NFR-05", "The system must remain functional for certificate creation even when blockchain connectivity is unavailable.", "Reliability"),
    ("NFR-06", "The user interface must be operable in French and must function on standard web browsers without requiring any plugins.", "Usability"),
    ("NFR-07", "The system must support SQLite for development and PostgreSQL for production without code changes.", "Portability"),
    ("NFR-08", "All API endpoints must return appropriate HTTP status codes and JSON error messages for failed requests.", "Maintainability"),
]

for i, (nid, req, cat) in enumerate(nfr, start=1):
    tbl2.rows[i].cells[0].text = nid
    tbl2.rows[i].cells[1].text = req
    tbl2.rows[i].cells[2].text = cat
    for j in range(3):
        for p in tbl2.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

add_fig(doc, "Table 4.2: Non-Functional Requirements of Certichain Africa")

add_heading2(doc, "4.2 System Design")

add_heading3(doc, "4.2.1 Use Case Diagram")

add_body(doc, "The Use Case Diagram for Certichain Africa identifies two primary actors and one external system actor. The first actor is the Institution Administrator, who represents any authenticated user of the platform acting on behalf of an educational institution. The second actor is the Public User, who represents any individual who visits the platform's public verification interface without logging in. The external system actor is the Blockchain Network, which automatically responds to smart contract queries during the verification process.")

add_body(doc, "The use cases for the Institution Administrator are as follows: Register Institution Account, Log In with OTP, Reset Password, View Dashboard, Create Certificate (with sub-cases: Select Certificate Type, Enter Recipient Details, Generate PDF, Anchor to Blockchain, Upload to IPFS), View Certificate List, Filter Certificates by Type and Status, Download Certificate PDF, Delete Certificate, Update Profile, Change Password, Update Wallet Address, Upgrade Subscription Plan via MoMo, Preview Certificate Template, and Download Certificate Template.")

add_body(doc, "The use cases for the Public User are as follows: Verify Certificate by File Upload, Verify Certificate by Blockchain Hash, and View Public Certificate Details via QR Code Link.")

add_body(doc, "The Blockchain Network actor interacts with the system during the Certificate Anchoring use case (triggered by the Institution Administrator) and the Verify Certificate on Blockchain use case (triggered by the Public User or the platform during the verification process).")

add_fig(doc, "Figure 4.1: Use Case Diagram for Certichain Africa - to be drawn in draw.io or Lucidchart using the description above")

add_heading3(doc, "4.2.2 Class Diagram")

add_body(doc, "The Class Diagram for Certichain Africa reflects the three SQLAlchemy model classes defined in models.py, together with their attributes, methods, and relationships.")

add_body(doc, "The Institution class has the following attributes: id (Integer, primary key), name (String 255, not null), email (String 255, unique, not null, indexed), password_hash (String 255, not null), is_verified (Boolean, default True), created_at (DateTime), updated_at (DateTime), wallet_address (String 255, nullable), reset_token (String 128, nullable, indexed), reset_token_expiry (DateTime, nullable), plan (String 50, default free), and plan_expires_at (DateTime, nullable). Its methods are set_password(password), which hashes and stores the password, check_password(password), which verifies a password against the stored hash, cert_quota(), which returns the number of certificates issued this month versus the institution's plan limit, and to_dict(), which serialises the object to a JSON-compatible dictionary.")

add_body(doc, "The Certificate class has the following attributes: id (Integer, primary key), institution_id (Integer, foreign key to institutions.id, not null), certificate_type (String 50, not null), recipient_name (String 255, not null), recipient_email (String 255, nullable), domain (String 255, nullable), mention (String 100, nullable), data (JSON, nullable), file_hash (String 255, nullable), ipfs_hash (String 255, nullable), blockchain_hash (String 255, nullable), status (String 50, default created), created_at (DateTime), and updated_at (DateTime). Its method is to_dict(), which serialises the certificate to a JSON-compatible dictionary.")

add_body(doc, "The Payment class has the following attributes: id (Integer, primary key), institution_id (Integer, foreign key to institutions.id, not null), reference_id (String 36, unique, not null), plan (String 50, not null), billing (String 20, default monthly), amount (Integer, not null), currency (String 10, default XAF), phone (String 20, not null), status (String 20, default pending), created_at (DateTime), and updated_at (DateTime). Its method is to_dict(), which serialises the payment to a JSON-compatible dictionary.")

add_body(doc, "The relationships are as follows: Institution has a one-to-many relationship with Certificate (one institution can have many certificates), and Institution has a one-to-many relationship with Payment (one institution can have many payments). Both relationships are configured with cascade all, delete-orphan, meaning that when an institution is deleted, all its associated certificates and payments are also deleted.")

add_fig(doc, "Figure 4.2: Class Diagram for Certichain Africa - to be drawn using the attribute and relationship descriptions above")

add_heading3(doc, "4.2.3 Entity-Relationship Diagram")

add_body(doc, "The Entity-Relationship (ER) Diagram represents the database schema of Certichain Africa. It consists of three entities: institutions, certificates, and payments.")

add_body(doc, "The institutions entity has the following attributes: id (PK, INT), name (VARCHAR 255, NOT NULL), email (VARCHAR 255, UNIQUE, NOT NULL), password_hash (VARCHAR 255, NOT NULL), is_verified (BOOLEAN), created_at (DATETIME), updated_at (DATETIME), wallet_address (VARCHAR 255), reset_token (VARCHAR 128), reset_token_expiry (DATETIME), plan (VARCHAR 50), and plan_expires_at (DATETIME).")

add_body(doc, "The certificates entity has the following attributes: id (PK, INT), institution_id (FK, INT, references institutions.id), certificate_type (VARCHAR 50, NOT NULL), recipient_name (VARCHAR 255, NOT NULL), recipient_email (VARCHAR 255), domain (VARCHAR 255), mention (VARCHAR 100), data (JSON), file_hash (VARCHAR 255), ipfs_hash (VARCHAR 255), blockchain_hash (VARCHAR 255), status (VARCHAR 50), created_at (DATETIME), and updated_at (DATETIME).")

add_body(doc, "The payments entity has the following attributes: id (PK, INT), institution_id (FK, INT, references institutions.id), reference_id (VARCHAR 36, UNIQUE, NOT NULL), plan (VARCHAR 50, NOT NULL), billing (VARCHAR 20), amount (INT, NOT NULL), currency (VARCHAR 10), phone (VARCHAR 20, NOT NULL), status (VARCHAR 20), created_at (DATETIME), and updated_at (DATETIME).")

add_body(doc, "The relationships are: institutions to certificates is one-to-many (one institution issues many certificates), and institutions to payments is one-to-many (one institution makes many payments). Both foreign key columns have an ON DELETE CASCADE constraint, implemented through SQLAlchemy's cascade option.")

add_fig(doc, "Figure 4.3: Entity-Relationship Diagram - to be drawn using the entity and relationship descriptions above")

add_heading3(doc, "4.2.4 System Architecture Diagram")

add_body(doc, "The system architecture of Certichain Africa follows the MVC pattern extended with a set of external service integrations. The architecture has four layers.")

add_body(doc, "The Presentation Layer consists of Jinja2 HTML templates rendered on the server side and delivered to the user's browser. Dynamic operations, such as filtering a list of certificates, are handled by JavaScript fetch calls to the API layer. The templates are stored in the templates directory and include pages for landing, signup, login, OTP verification, dashboard, certificate creation, certificate management, public verification, templates library, settings, and pricing.")

add_body(doc, "The Application Layer consists of the Flask route functions defined in app.py. These functions receive HTTP requests, validate inputs, coordinate with the model layer to read or write data, interact with external services as needed, and return responses either as rendered HTML pages or as JSON API responses.")

add_body(doc, "The Data Layer consists of the SQLAlchemy models defined in models.py and the underlying database, which is SQLite during development and can be configured as PostgreSQL for production through the DATABASE_URL environment variable. The data layer also includes the local file system storage for certificate PDFs, stored in the certs/uploads directory.")

add_body(doc, "The External Services Layer consists of four external APIs that the application interacts with: the Polygon blockchain node (accessed through an Infura RPC endpoint), the Pinata IPFS cloud service, the SMTP email server (configured through Flask-Mail), and the MTN MoMo Collection API.")

add_fig(doc, "Figure 4.4: System Architecture Diagram (MVC + External Services) - to be drawn using the layer descriptions above")

add_heading2(doc, "4.3 Implementation")

add_heading3(doc, "4.3.1 Technology Stack Rationale")

add_body(doc, "The technology choices described in Chapter 3 were validated during the implementation phase. Flask's routing system and extension architecture proved well-suited to the project's requirements. The ability to use Flask extensions selectively, adding SQLAlchemy for database access, Flask-Mail for email, and Flask-CORS for API endpoints without loading everything by default, kept the application lightweight and easy to understand. Web3.py version 6 introduced breaking changes from version 5, including new import paths for middleware and updated function naming conventions, and the implementation handled these differences carefully to ensure compatibility.")

add_body(doc, "ReportLab's canvas API provided the level of graphical control needed to produce the professional certificate designs. The ability to draw arbitrary shapes, position text at precise coordinates, apply custom colours, and embed QR code images allowed the four certificate templates to be implemented entirely in Python without any dependency on image editing software or external design assets.")

add_heading3(doc, "4.3.2 Feature 1: Institution Registration and Authentication")

add_body(doc, "The authentication system is one of the most critical components of Certichain Africa. An institution that cannot securely authenticate its staff cannot protect the certificate issuance authority that the platform grants it. The authentication system was therefore designed with multiple layers of protection.")

add_body(doc, "The institution registration route at /signup validates that the submitted email address is unique and conforms to a standard email format, that the password is at least eight characters long, contains at least one uppercase letter, at least one lowercase letter, and at least one digit, and that the password and password confirmation fields match. If all validations pass, a new Institution record is created and its password is stored as a PBKDF2-SHA256 hash using Werkzeug's generate_password_hash function.")

add_code(doc, "app.py (signup validation)", """def validate_password(password):
    if len(password) < 8:
        return False, "Le mot de passe doit contenir au moins 8 caracteres"
    if not re.search(r'[A-Z]', password):
        return False, "Le mot de passe doit contenir au moins une majuscule"
    if not re.search(r'[a-z]', password):
        return False, "Le mot de passe doit contenir au moins une minuscule"
    if not re.search(r'[0-9]', password):
        return False, "Le mot de passe doit contenir au moins un chiffre"
    return True, "Mot de passe valide"

institution = Institution(name=name, email=email, is_verified=True)
institution.set_password(password)
db.session.add(institution)
db.session.commit()""")

add_body(doc, "The login route at /login handles two different scenarios. When Flask-Mail is configured with valid SMTP credentials, the route validates the institution's email and password, then generates a six-digit OTP code, stores its SHA-256 hash in the session along with an expiry timestamp and an attempts counter, sends the code to the institution's email address, and redirects the user to the OTP verification page. When Flask-Mail is not configured, the route skips the OTP step and creates the session directly, allowing the platform to function without an SMTP server during development or testing.")

add_code(doc, "app.py (OTP generation and storage)", """OTP_MAX_SEND_PER_HOUR = 5
OTP_RESEND_COOLDOWN   = 60
OTP_MAX_ATTEMPTS      = 5
OTP_EXPIRY_MINUTES    = 10

def _generate_otp() -> str:
    return ''.join(secrets.choice('0123456789') for _ in range(6))

def _store_otp(email, institution_id, code):
    now = datetime.now().timestamp()
    session['otp_pending_id']    = institution_id
    session['otp_pending_email'] = email
    session['otp_hash']          = _hash_otp(code)
    session['otp_expiry']        = now + OTP_EXPIRY_MINUTES * 60
    session['otp_attempts']      = 0
    session['otp_last_sent']     = now""")

add_body(doc, "The OTP verification route at /login/verify-otp checks that the session contains a pending OTP, that the OTP has not expired, and that the number of failed attempts has not reached the maximum of five. It compares the SHA-256 hash of the entered code with the stored hash. If the hashes match, the OTP session data is cleared, a proper authenticated session is created, and the user is redirected to the dashboard. If the hashes do not match, the attempt counter is incremented and the user is informed of how many attempts remain. When all five attempts are exhausted, the session is invalidated and the user must restart the login process.")

add_body(doc, "An in-memory rate limiter prevents a single email address from requesting more than five OTP codes per hour. This protects against automated attacks that might attempt to exhaust the OTP code space by requesting a large number of codes.")

add_fig(doc, "Figure 4.6: Authentication Flow Sequence Diagram - to be drawn showing: Browser -> Flask /login -> validate creds -> generate OTP -> Flask-Mail -> SMTP -> email inbox -> Browser -> Flask /verify-otp -> compare hash -> create session -> redirect to dashboard")

add_heading3(doc, "4.3.3 Feature 2: Password Reset")

add_body(doc, "The password reset flow is initiated at /forgot-password, where the institution enters its email address. The system responds with a success message regardless of whether the email is found in the database, preventing email enumeration attacks where an attacker could determine which email addresses have registered accounts by observing the response.")

add_body(doc, "If the email is found and Flask-Mail is configured, the system generates a cryptographically random 48-byte URL-safe token using Python's secrets.token_urlsafe function, stores the token and a one-hour expiry timestamp on the Institution record, and sends an HTML email containing a password reset link to the institution's email address. The reset link leads to the /reset-password/<token> route, which validates the token and its expiry before allowing the password to be changed.")

add_code(doc, "app.py (password reset token generation)", """token  = secrets.token_urlsafe(48)
expiry = datetime.now() + timedelta(hours=1)
institution.reset_token        = token
institution.reset_token_expiry = expiry
db.session.commit()

base = os.getenv('APP_BASE_URL', request.host_url.rstrip('/'))
reset_url = f"{base}/reset-password/{token}" """)

add_heading3(doc, "4.3.4 Feature 3: Certificate Creation Pipeline")

add_body(doc, "The certificate creation pipeline is the most technically complex feature of the platform. It is implemented primarily in the /api/certificates/create route in app.py, with the PDF generation delegated to pdf_generator.py. The pipeline consists of seven steps that are executed in sequence for each certificate.")

add_body(doc, "Step 1 is data validation and database record creation. The route receives a JSON payload from the frontend containing the certificate type, recipient name, recipient email, domain, mention, and other certificate-specific fields. A new Certificate record is created in the database with status set to created. This record is saved immediately so that it has a database ID that can be used in subsequent steps.")

add_body(doc, "Step 2 is PDF generation. The appropriate PDF generation function is selected based on the certificate type: create_diploma_pdf for diplomas, create_certification_pdf for certifications, create_badge_pdf for badges, and create_cyber_diploma_pdf for the cybersecurity template. The selected function is called with a dictionary containing all the certificate data, including the institution name, recipient name, domain, mention, graduation date, certificate number, and placeholder values for the blockchain hash and verification URL. The function returns a BytesIO buffer containing the PDF data.")

add_body(doc, "Step 3 is initial file saving. The PDF buffer is written to the certs/uploads directory with the filename cert_{id}.pdf, where {id} is the database ID of the certificate record. This file is used in subsequent steps for hashing and IPFS upload.")

add_body(doc, "Step 4 is blockchain anchoring. The SHA-256 hash of the saved PDF file is computed, and the Ethereum keccak256 hash of that SHA-256 hash is used as the unique certificate identifier for the smart contract call. If the system is blockchain-configured, the issueCertificate function is called on the deployed CertiChain contract, passing the certificate identifier, an empty IPFS hash placeholder, and the recipient's name. The resulting transaction hash is stored as the certificate's blockchain_hash.")

add_code(doc, "app.py (blockchain anchoring)", """file_hash = generate_file_hash(file_path)

if BLOCKCHAIN_CONFIGURED:
    cert_id_bytes = w3.solidity_keccak(['string'], [file_hash])
    nonce = w3.eth.get_transaction_count(CHECKED_ISSUER)
    tx = contract.functions.issueCertificate(
        cert_id_bytes, '', cert.recipient_name
    ).build_transaction({
        'chainId': 137,
        'gas': 500000,
        'gasPrice': w3.eth.gas_price,
        'nonce': nonce,
    })
    signed  = w3.eth.account.sign_transaction(tx, private_key=ISSUER_PRIVATE_KEY)
    tx_hash = w3.eth.send_raw_transaction(signed.raw_transaction)
    w3.eth.wait_for_transaction_receipt(tx_hash)
    blockchain_hash = w3.to_hex(tx_hash)
else:
    blockchain_hash = 'sha256:' + file_hash""")

add_body(doc, "Step 5 is PDF regeneration. Once the blockchain hash is known, whether it is an on-chain transaction hash or a local SHA-256 fingerprint, the PDF generation function is called a second time with the actual blockchain hash included in the data payload. This ensures that the final PDF file contains the real blockchain hash printed at the bottom, not a placeholder. The regenerated PDF overwrites the initial file in the certs/uploads directory.")

add_body(doc, "Step 6 is IPFS upload. The final PDF file is uploaded to Pinata IPFS using the upload_to_ipfs function, which sends an HTTP POST request to the Pinata API endpoint with the file as a multipart attachment. If the upload succeeds, the returned IPFS Content Identifier is stored as the certificate's ipfs_hash. If the upload fails, a local fallback identifier of the form local:{id} is stored instead.")

add_code(doc, "app.py (IPFS upload via Pinata)", """def upload_to_ipfs(file_path):
    url = "https://api.pinata.cloud/pinning/pinFileToIPFS"
    headers = {
        "pinata_api_key": PINATA_API_KEY,
        "pinata_secret_api_key": PINATA_SECRET_KEY
    }
    with open(file_path, "rb") as f:
        response = requests.post(url, files={"file": f}, headers=headers)
    if response.status_code == 200:
        return response.json()["IpfsHash"]
    raise Exception(f"IPFS Error: {response.text}")""")

add_body(doc, "Step 7 is final hash computation and database commit. The SHA-256 hash of the final PDF file, the one that will be downloaded by the institution and the one whose hash is printed on the certificate, is computed and stored as the certificate's file_hash. This is the hash that the public verification system will use to match an uploaded file against the database. The certificate status is set to issued, and the database session is committed.")

add_fig(doc, "Figure 4.5: Certificate Issuance Sequence Diagram - to be drawn showing: Institution Browser -> Flask /api/certificates/create -> validate -> DB insert -> pdf_generator.py -> save PDF -> Web3.py -> Polygon -> store tx hash -> regenerate PDF -> Pinata IPFS -> store CID -> recompute hash -> DB commit -> return JSON")

add_heading3(doc, "4.3.5 Feature 4: PDF Generation Engine")

add_body(doc, "The PDF generation engine is implemented in pdf_generator.py and consists of four distinct functions, one for each certificate template. Each function creates a ReportLab canvas on an in-memory BytesIO buffer, draws the certificate content using the canvas API, saves the canvas, and returns the buffer.")

add_body(doc, "The diploma template (create_diploma_pdf) produces a landscape A4 document with a white background featuring a central light blue panel, decorative corner circles in teal and navy, and gold accent lines at the top and bottom. The institution name is printed in bold navy at the top, followed by the phrase DECERNE FIEREMENT in large spaced uppercase letters, the recipient's name in a large bold font with a navy underline, a description of the programme completed, the domain or specialisation in gold, the graduation date, an optional mention, and two signature blocks flanking a decorative circular seal at the bottom.")

add_body(doc, "The certification template (create_certification_pdf) produces a portrait A4 document with a dark navy background featuring scattered star dots in white and gold, a central mid-navy panel, gold horizontal bars at the top and bottom, and the words CERTIFICATE OF ACHIEVEMENT in large white text at the top. The recipient's name appears in white below the header with a gold underline, followed by a description, the domain in gold, and optional competency pills rendered as small rounded boxes arranged in a grid.")

add_code(doc, "pdf_generator.py (diploma template excerpt)", """def create_diploma_pdf(data=None):
    if data is None:
        data = {}
    recipient = data.get('recipient_name', 'Jean-Baptiste Kouassi')
    domain    = data.get('domain', 'Communication and Management')

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    W, H = landscape(A4)

    c.setFillColor(white)
    c.rect(0, 0, W, H, fill=1, stroke=0)

    c.setFillColor(NAVY)
    c.setFont("Helvetica-Bold", 30)
    c.drawCentredString(W/2, H - 5.8*cm, recipient)

    c.setFillColor(GOLD)
    c.setFont("Helvetica-Bold", 15)
    c.drawCentredString(W/2, H - 8.3*cm, domain)

    c.save()
    buffer.seek(0)
    return buffer""")

add_body(doc, "The badge template (create_badge_pdf) produces a portrait A4 document with a white background and a navy header band at the top. At the boundary between the header and the body, a layered circular medallion is drawn using concentric circles in gold, white, and navy. The body contains the recipient name with a gold underline, the competence area in gold, and an information grid with three cells showing the level, validity period, and issue date. A blockchain verification box and signature blocks complete the layout.")

add_body(doc, "The cyber diploma template (create_cyber_diploma_pdf) produces a landscape A4 document with a very dark near-black background, simulating the visual aesthetic of cybersecurity and technology. The decorative elements are magenta hexagonal clusters drawn in the top-right and bottom-left corners, with magenta neon accent lines at the top and bottom edges. The text is white and magenta on the dark background, creating a high-contrast, professional look appropriate for technology credentials.")

add_body(doc, "Each template generates a QR code using the qrcode[pil] library. The QR code encodes the verification URL for the certificate, which follows the pattern {base_url}/verify/{cert_id}. The QR code image is rendered as a PNG in memory and drawn onto the PDF canvas at a fixed position in the bottom-right corner. A Scan to verify label is printed below the QR code.")

add_heading3(doc, "4.3.6 Feature 5: Smart Contract and Blockchain Integration")

add_body(doc, "The Certichain Africa smart contract is written in Solidity and deployed on the Polygon network. The contract is named CertiChain and is defined in CertiChain.sol. It contains a mapping from a bytes32 certificate identifier to a Certificate struct, which stores the IPFS hash, recipient name, issue timestamp, and an existence flag.")

add_code(doc, "CertiChain.sol", """// SPDX-License-Identifier: MIT
pragma solidity ^0.8.0;

contract CertiChain {
    struct Certificate {
        string ipfsHash;
        string recipientName;
        uint256 issueDate;
        bool exists;
    }

    mapping(bytes32 => Certificate) public certificates;
    address public owner;

    modifier onlyOwner() {
        require(msg.sender == owner, "Not authorized");
        _;
    }

    function issueCertificate(
        bytes32 _certificateId,
        string memory _ipfsHash,
        string memory _recipientName
    ) public onlyOwner {
        require(!certificates[_certificateId].exists, "Already exists");
        certificates[_certificateId] = Certificate(
            _ipfsHash, _recipientName, block.timestamp, true
        );
    }

    function verifyCertificate(bytes32 _certificateId)
        public view returns (bool) {
        return certificates[_certificateId].exists;
    }
}""")

add_body(doc, "The contract's onlyOwner modifier ensures that only the wallet that deployed the contract, the platform's issuer wallet, can call issueCertificate. This prevents any external party from writing a fraudulent certificate record to the contract's mapping. The verifyCertificate function is a view function that does not require any gas to call, which means anyone can check whether a certificate exists on-chain at no cost.")

add_body(doc, "On the Flask application side, the contract is loaded at startup by reading the contract address from the CONTRACT_ADDRESS environment variable and the contract ABI from the contract_abi.json file. A Web3 connection is established to the Polygon RPC endpoint, and a contract object is created using w3.eth.contract(address=CHECKED_CONTRACT, abi=CONTRACT_ABI). The BLOCKCHAIN_CONFIGURED flag is set to True only if all required environment variables are present and the contract address is valid, ensuring that the application does not crash at startup when blockchain credentials are missing.")

add_heading3(doc, "4.3.7 Feature 6: Certificate Verification")

add_body(doc, "The certificate verification system provides two paths for verifying a certificate's authenticity. The primary path is file-based verification, available at the /verify route. A user uploads the PDF file of the certificate they want to verify. The system saves the file, computes its SHA-256 hash, and looks for a matching record in the Certificate table's file_hash column. If a record is found, the system returns the certificate details and confirms its authenticity. If no record is found in the database, the system falls back to checking the Polygon blockchain by deriving the keccak256 hash of the file hash and calling verifyCertificate on the smart contract.")

add_code(doc, "app.py (verification by file hash)", """file_hash = generate_file_hash(file_path)
cert = Certificate.query.filter_by(file_hash=file_hash).first()

if cert:
    on_chain = cert.blockchain_hash and cert.blockchain_hash.startswith('0x')
    return jsonify({
        "verified": True,
        "recipient": cert.recipient_name,
        "certificate_type": cert.certificate_type,
        "domain": cert.domain,
        "on_chain": on_chain,
        "message": "Certificat authentique"
    })

if contract:
    cert_id = w3.solidity_keccak(['string'], [file_hash])
    exists  = contract.functions.verifyCertificate(cert_id).call()
    if exists:
        data = contract.functions.getCertificate(cert_id).call()
        return jsonify({"verified": True, "recipient": data[1]})

return jsonify({"verified": False, "message": "Certificat non trouve"})""")

add_body(doc, "The secondary verification path is hash-based, available at the /verify-hash endpoint. A user or application submits a blockchain hash or SHA-256 hash as a JSON body parameter, and the system searches the Certificate table's blockchain_hash column for a matching record. This path is particularly useful for automated integrations where a system wants to verify a certificate without downloading and uploading the PDF file.")

add_body(doc, "A third verification path is the public certificate page, available at /verify/{cert_id}. This is the URL encoded in the QR code on each generated certificate. When a user scans the QR code with a smartphone, they are taken to this page, which displays the certificate details in a clean, readable format without requiring any file upload. The page shows the recipient name, certificate type, domain, issuing institution, issue date, and a confirmation that the certificate is either on-chain or locally hash-verified.")

add_heading3(doc, "4.3.8 Feature 7: Certificate Management")

add_body(doc, "The certificate management interface allows authenticated institution administrators to view all the certificates they have issued, filter them by type and status, search by recipient name, download individual certificates as PDF files, and delete certificates they no longer want to maintain.")

add_body(doc, "The /api/certificates endpoint retrieves all certificates belonging to the logged-in institution from the database and returns them as a JSON array. Each certificate object in the array includes all the fields defined in the Certificate model's to_dict() method. The frontend JavaScript function loadCertificates calls this endpoint when the certificates page loads and stores the results in a local array. Filtering is performed on the client side by the filterCertificates function, which applies the selected type, status, and search filters to the local array without making additional API calls.")

add_code(doc, "templates/certificates.html (JavaScript certificate filtering)", """function filterCertificates() {
    const type   = document.getElementById('filterType').value;
    const status = document.getElementById('filterStatus').value;
    const search = document.getElementById('searchInput').value.toLowerCase();

    const filtered = allCertificates.filter(cert => {
        const typeMatch   = !type   || cert.certificate_type === type;
        const statusMatch = !status || cert.status === status;
        const searchMatch = !search || cert.recipient_name.toLowerCase().includes(search);
        return typeMatch && statusMatch && searchMatch;
    });
    renderCertificates(filtered);
}""")

add_body(doc, "Certificate download is handled by the /certificate/{cert_id}/download route. This route first checks whether the saved PDF file exists in the certs/uploads directory. If the file exists, it is served directly using Flask's send_file function. If the file does not exist (for example, if the server's storage was reset), the route regenerates the PDF from the certificate's stored data and saves it before serving it. This regeneration fallback ensures that download functionality remains available even when file storage is not persistent.")

add_body(doc, "Certificate deletion is handled by the /api/certificates/{cert_id} DELETE endpoint. The endpoint verifies that the certificate belongs to the requesting institution before deleting it, preventing one institution from deleting another's certificates. The deletion removes only the database record; the PDF file in the certs/uploads directory is not automatically deleted, which could be addressed in a future maintenance operation.")

add_heading3(doc, "4.3.9 Feature 8: Subscription and Mobile Money Payment")

add_body(doc, "The subscription system defines three plans for institutions: Free, Starter, and Academique. The Free plan allows up to five certificate issuances per calendar month at no cost. The Starter plan, priced at 19,900 XAF per month or 190,800 XAF per year, allows up to fifty certificates per month. The Academique plan, priced at 49,900 XAF per month or 478,800 XAF per year, allows up to three hundred certificates per month. An Enterprise tier is also defined in the code with an effectively unlimited quota, intended for future use with custom pricing.")

add_body(doc, "The payment flow begins when an institution clicks a subscribe button on the pricing page. The frontend sends a POST request to /api/payment/initiate with the selected plan, billing period, and the institution's MTN phone number. The route validates the inputs, normalises the phone number to the international format (2376XXXXXXXX), calculates the amount in XAF based on the plan and billing period, generates a UUID as the payment reference ID, and calls the momo_api.request_to_pay function.")

add_code(doc, "momo.py (MoMo payment request)", """def request_to_pay(phone, amount, currency, reference_id, note):
    token = _token()
    r = requests.post(
        f"{BASE_URL}/collection/v1_0/requesttopay",
        json={
            "amount": str(amount),
            "currency": currency,
            "externalId": reference_id,
            "payer": {"partyIdType": "MSISDN", "partyId": phone},
            "payerMessage": note,
            "payeeNote": note,
        },
        headers=_headers(token, reference_id),
        timeout=20,
    )
    return r.status_code == 202""")

add_body(doc, "The request_to_pay function first obtains a Bearer token from the MoMo API by sending the API user credentials as a Base64-encoded Basic authentication header. It then sends the payment request to the /collection/v1_0/requesttopay endpoint. If the endpoint returns HTTP 202 Accepted, the payment request has been successfully submitted and the MTN user will receive a prompt on their phone to approve the payment. The function returns True on success and False on failure.")

add_body(doc, "After the payment request is submitted, a Payment record is created in the database with status pending. The frontend polls the /api/payment/status/{reference_id} endpoint at regular intervals. Each poll calls momo_api.get_status, which queries the MoMo API for the current status of the payment. When the status changes to SUCCESSFUL, the institution's plan is upgraded to the selected plan and the plan_expires_at date is set to thirty days or three hundred and sixty-five days in the future depending on the billing period. When the status changes to FAILED, the payment record is marked as failed and no plan change is made.")

add_body(doc, "The MoMo webhook at /api/payment/callback provides an optional push notification path: MTN can call this endpoint directly when a payment status changes, allowing the platform to update the payment status without waiting for the next poll. In practice, the polling mechanism is more reliable because it does not depend on MTN's webhook delivery infrastructure.")

add_heading3(doc, "4.3.10 Feature 9: Settings Management")

add_body(doc, "The settings page allows institution administrators to manage three aspects of their account: their institution profile, their password, and their blockchain wallet address.")

add_body(doc, "Profile management is handled by the /update-profile route, which accepts a POST request with the updated institution name. The route validates that the name is not empty, updates the Institution record in the database, and flashes a success message using Flask's flash system.")

add_body(doc, "Password change is handled by the /change-password route, which requires the user to enter their current password, a new password, and a password confirmation. The route verifies the current password using check_password, validates the new password against the strength criteria, checks that the new password and confirmation match, and updates the password hash if all checks pass.")

add_body(doc, "Wallet address management is handled by the /update-wallet route, which accepts a POST request with an Ethereum wallet address. The route validates the address against the standard Ethereum address format using a regular expression requiring the 0x prefix followed by exactly forty hexadecimal characters, and stores the validated address in the Institution record's wallet_address field. This feature allows institutions to associate their own Ethereum wallet with their account, which is relevant for future features such as allowing institutions to issue certificates from their own wallets rather than the platform operator's wallet.")

add_heading3(doc, "4.3.11 Feature 10: Certificate Templates Library")

add_body(doc, "The templates library at /templates provides a visual showcase of the four available certificate designs. Each design is presented on the page with a description, and two action buttons allow the institution to preview the template as a PDF in the browser or download it as a template file.")

add_body(doc, "The preview route at /templates/preview/{template_type} generates a sample PDF using the appropriate generation function with placeholder data, sets the response MIME type to application/pdf, and sets the Content-Disposition header to inline, which instructs the browser to display the PDF in its built-in PDF viewer rather than downloading it. The download route at /templates/download/{template_type} does the same but sets Content-Disposition to attachment, which triggers a file download.")

add_code(doc, "app.py (template preview route)", """@app.route('/templates/preview/<template_type>')
def template_preview(template_type):
    from pdf_generator import (create_diploma_pdf, create_certification_pdf,
                               create_badge_pdf, create_cyber_diploma_pdf)
    pdf_map = {
        'diplome':       create_diploma_pdf,
        'certification': create_certification_pdf,
        'badge':         create_badge_pdf,
        'cyber':         create_cyber_diploma_pdf
    }
    if template_type not in pdf_map:
        return render_template('404.html'), 404
    pdf_buffer = pdf_map[template_type]()
    return send_file(pdf_buffer, mimetype='application/pdf',
                     as_attachment=False,
                     download_name=f'apercu_{template_type}.pdf')""")

add_heading2(doc, "4.4 Quality Assurance and Testing")

add_heading3(doc, "4.4.1 Testing Approach")

add_body(doc, "The testing approach for Certichain Africa was primarily manual functional testing, supplemented by targeted scenario-based testing for the security-critical components. A full automated test suite was not developed within the scope of this project, which is acknowledged as a limitation and identified as a priority for future development. However, each feature was tested thoroughly through its complete user journey before being considered complete.")

add_heading3(doc, "4.4.2 Authentication Testing")

add_body(doc, "The authentication system was tested through a set of scenarios designed to verify both normal operation and edge cases. In the normal registration scenario, a new institution was registered with valid credentials and the system was verified to create a database record with a hashed password and to redirect to the login page with a success message. In the duplicate email scenario, attempting to register with an email already in the database was verified to display an appropriate error message without creating a duplicate record.")

add_body(doc, "In the OTP flow scenario, a login with valid credentials was verified to generate an OTP, send it to the configured email address, and redirect to the OTP page. Entering the correct OTP within the time limit was verified to create a session and redirect to the dashboard. Entering an incorrect OTP was verified to decrement the attempts counter and display the remaining attempts. Entering an incorrect OTP five times was verified to clear the session and redirect to the login page with an appropriate error message.")

add_body(doc, "The OTP expiry was tested by temporarily reducing the OTP_EXPIRY_MINUTES constant to one minute and waiting for the code to expire before attempting to verify it. The system correctly rejected the expired code and prompted the user to log in again.")

add_heading3(doc, "4.4.3 Certificate Creation Testing")

add_body(doc, "Certificate creation was tested for each of the four template types. For each type, a certificate was created with complete form data and the resulting PDF was downloaded and inspected visually to verify that all fields were correctly populated, that the QR code was present and scannable, and that the blockchain hash was printed on the certificate. The QR codes were scanned using a smartphone camera application and verified to resolve to the correct public verification URL.")

add_body(doc, "The blockchain fallback was tested by temporarily setting the ISSUER_ADDRESS environment variable to an empty string, which disables blockchain configuration. A certificate was then created and verified to be assigned a sha256:-prefixed hash in the blockchain_hash field rather than a 0x-prefixed transaction hash. The certificate was then downloaded and verified to display the sha256 hash on the PDF.")

add_body(doc, "The IPFS fallback was tested by temporarily using an invalid Pinata API key. A certificate created in this state was verified to receive a local:{id} IPFS hash and to still be downloadable and verifiable through the database path.")

add_heading3(doc, "4.4.4 Verification Testing")

add_body(doc, "The file-based verification was tested by uploading a certificate PDF that had been created through the platform and verifying that the system returned a positive verification result with the correct certificate details. A modified version of the same PDF, produced by opening it in a PDF editor and making a small change, was then uploaded and verified to return a negative result, confirming that the SHA-256 hash comparison correctly detects tampering.")

add_body(doc, "The hash-based verification was tested by submitting the blockchain hash of an issued certificate to the /verify-hash endpoint and verifying that the system returned the certificate details. Submitting a randomly generated hash was verified to return a not-found response.")

add_heading3(doc, "4.4.5 Payment Testing")

add_body(doc, "The MoMo payment integration was tested using the MTN MoMo Sandbox environment. A payment request was initiated using a test phone number provided by the sandbox documentation, and the payment status polling was verified to update the payment record and the institution's plan when the sandbox reported a SUCCESSFUL status.")

add_heading3(doc, "4.4.6 Test Summary")

tbl = doc.add_table(rows=11, cols=4)
tbl.style = 'Table Grid'
for i, h in enumerate(["Test ID", "Feature", "Scenario", "Result"]):
    tbl.rows[0].cells[i].text = h
    for p in tbl.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

tests = [
    ("TC-01", "Registration", "Valid new institution registration", "Pass"),
    ("TC-02", "Registration", "Duplicate email rejected", "Pass"),
    ("TC-03", "Login + OTP", "Correct OTP accepted within time limit", "Pass"),
    ("TC-04", "Login + OTP", "5 wrong OTPs locks session", "Pass"),
    ("TC-05", "Login + OTP", "Expired OTP rejected", "Pass"),
    ("TC-06", "Certificate", "Diploma PDF generated with all fields", "Pass"),
    ("TC-07", "Certificate", "QR code resolves to correct URL", "Pass"),
    ("TC-08", "Certificate", "Blockchain fallback on missing credentials", "Pass"),
    ("TC-09", "Verification", "Genuine certificate file verified", "Pass"),
    ("TC-10", "Verification", "Tampered certificate file rejected", "Pass"),
]

for i, (tid, feat, scen, res) in enumerate(tests, start=1):
    tbl.rows[i].cells[0].text = tid
    tbl.rows[i].cells[1].text = feat
    tbl.rows[i].cells[2].text = scen
    tbl.rows[i].cells[3].text = res
    for j in range(4):
        for p in tbl.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

add_fig(doc, "Table 4.9: Test Cases and Results")

add_heading2(doc, "4.5 Application Overview")

add_body(doc, "The following section provides a screen-by-screen walkthrough of the completed application. Screenshot placeholders are provided for each key interface, to be replaced with actual screenshots before the final submission.")

add_heading3(doc, "4.5.1 Landing Page")
add_body(doc, "The landing page at / is the first point of contact for a new visitor. It presents an overview of the platform's features, a call-to-action to sign up, and links to the pricing and verification pages. The page is designed to communicate the core value proposition quickly: that Certichain Africa enables institutions to issue certificates that cannot be forged.")
add_fig(doc, "Figure 4.7: Landing Page Screenshot - to be inserted")

add_heading3(doc, "4.5.2 Login and OTP Pages")
add_body(doc, "The login page at /login presents a simple form with email and password fields and a link to the password reset page. After successful credential validation, the user is redirected to the OTP verification page, which prompts them to enter the six-digit code sent to their email. The OTP page displays the masked email address, the number of attempts remaining, and a resend button that becomes active after the cooldown period expires.")
add_fig(doc, "Figure 4.8: Login Page Screenshot - to be inserted")
add_fig(doc, "Figure 4.9: OTP Verification Page Screenshot - to be inserted")

add_heading3(doc, "4.5.3 Dashboard")
add_body(doc, "The dashboard at /dashboard is the main control centre for an institution administrator. It displays a summary of the institution's certificate statistics, including the total number of certificates issued, the current plan and quota usage, and quick-access buttons to create a new certificate and view all certificates.")
add_fig(doc, "Figure 4.10: Institution Dashboard Screenshot - to be inserted")

add_heading3(doc, "4.5.4 Certificate Creation")
add_body(doc, "The certificate creation page at /create-certificate presents a multi-field form for entering the certificate details. The form includes a dropdown for selecting the certificate type, text fields for the recipient name and email, and additional fields for domain, mention, graduation date, and competencies. When the form is submitted, a loading indicator is shown while the system generates the PDF and anchors the certificate.")
add_fig(doc, "Figure 4.11: Certificate Creation Form Screenshot - to be inserted")
add_fig(doc, "Figure 4.12: Generated Diploma PDF Sample - to be inserted")

add_heading3(doc, "4.5.5 Certificate Management")
add_body(doc, "The certificate management page at /my-certificates displays all certificates issued by the institution in a card-based grid layout. Each card shows the certificate type, recipient name, domain, status, creation date, and partial blockchain hash. Three action buttons on each card allow the administrator to view the certificate details, download the PDF, or delete the certificate. Filter dropdowns and a search input at the top of the page allow the list to be filtered by type, status, and recipient name.")
add_fig(doc, "Figure 4.13: Certificate Management Page Screenshot - to be inserted")

add_heading3(doc, "4.5.6 Public Verification Page")
add_body(doc, "The public verification page is accessible at /verify/{cert_id} and is the destination URL encoded in each certificate's QR code. It requires no login and displays the certificate details in a clean, trust-signalling layout. A green checkmark icon and the text Certificat Authentique are shown when the certificate is found. If the certificate is anchored on the blockchain, a blockchain badge is shown along with a truncated version of the transaction hash.")
add_fig(doc, "Figure 4.14: Public Verification Page Screenshot - to be inserted")

add_heading3(doc, "4.5.7 Pricing Page")
add_body(doc, "The pricing page at /pricing presents the three subscription plans in a side-by-side comparison layout. Each plan card shows the monthly and annual price in XAF, the number of certificates allowed per month, and a list of included features. A subscribe button on each paid plan card initiates the MoMo payment flow.")
add_fig(doc, "Figure 4.15: Pricing and Subscription Page Screenshot - to be inserted")

add_heading3(doc, "4.5.8 Settings Page")
add_body(doc, "The settings page at /settings is divided into three sections. The Profile section allows the institution name to be updated. The Security section allows the password to be changed by entering the current password and a new password twice. The Blockchain section allows an Ethereum wallet address to be associated with the institution account.")
add_fig(doc, "Figure 4.16: Settings Page Screenshot - to be inserted")

page_break(doc)

# ==============================================================================
# CHAPTER 4 PART 2: DISCUSSION
# ==============================================================================

add_heading1(doc, "CHAPTER 4 (PART 2)\nDISCUSSION")

add_heading2(doc, "4.6 Interpretation of Results")

add_body(doc, "The implementation of Certichain Africa produced a working web application that meets all fourteen of the defined functional requirements and all eight of the non-functional requirements to a satisfactory degree. This section interprets what those results mean in the context of the problem the project set out to address.")

add_body(doc, "The most important result is that the certificate creation pipeline works end-to-end. An authenticated institution administrator can log in, fill out a certificate form, and receive a professionally designed PDF within seconds. That PDF is self-contained in the sense that it carries, printed on its face, the cryptographic hash that anchors it to either the blockchain or the local verification system. A verifier who receives this certificate can scan the QR code with their smartphone, be taken to the public verification page, and see an immediate confirmation of authenticity without any additional tools or accounts. This end-to-end workflow is the core of the platform's value proposition, and it works.")

add_body(doc, "The dual anchoring strategy also produced an important result: the platform is resilient to blockchain infrastructure failures. During testing, when the blockchain was deliberately disabled, the platform continued to create certificates and assign them SHA-256 fingerprints. These fingerprints are weaker than blockchain records in that they do not provide the same level of independence from the platform operator, but they are sufficient for basic verification and they ensure that the certificate creation workflow is never blocked by external service failures. This resilience is particularly important for an African deployment context where network connectivity may be intermittent.")

add_body(doc, "The OTP two-factor authentication system worked correctly in all test scenarios. The rate limiting, attempt counting, expiry checking, and session invalidation all functioned as designed. The branded OTP email, rendered in HTML with the Certichain Africa visual identity, presents a professional appearance that builds trust with the receiving institution administrator. The password reset flow, which uses a cryptographically random token with a one-hour expiry, provides a secure and user-friendly path to account recovery.")

add_body(doc, "The MTN Mobile Money payment integration worked in the sandbox environment. The payment request API call, the status polling mechanism, and the plan upgrade logic all functioned correctly. This is a particularly meaningful result because mobile money API integration is not commonly documented in the open-source Python web development community, and the implementation developed for this project may serve as a useful reference for other developers in the region.")

add_body(doc, "The PDF generation engine produced high-quality certificate documents that are visually comparable to those produced by commercial platforms such as Accredible and Credly. The use of a custom colour palette, geometric decorations, and carefully chosen typography gives the certificates a professional appearance. The ability to generate four distinct templates from the same codebase, with full control over every visual element, demonstrates the power and flexibility of ReportLab as a PDF generation tool.")

add_heading2(doc, "4.7 Comparison with Existing Works")

add_body(doc, "This section compares Certichain Africa with the existing platforms analysed in the literature review, evaluating each on the dimensions that matter most for the African educational context.")

add_heading3(doc, "4.7.1 Comparison with Blockcerts")

add_body(doc, "Blockcerts and Certichain Africa share the core concept of anchoring academic credentials to a public blockchain. Both systems use a hash-based approach to link a credential document to a blockchain record, and both provide public verification without requiring the verifier to have an account on the issuing platform. The key differences are in accessibility and completeness.")

add_body(doc, "Blockcerts is a developer toolkit, not a ready-to-use platform. An institution wishing to use Blockcerts must set up its own server, deploy the Blockcerts issuer application, manage its own Bitcoin or Ethereum wallet, and design and generate certificate documents separately using other tools. This requires a level of technical expertise that is beyond the capacity of most African educational institutions. Certichain Africa, by contrast, is a fully hosted web application that can be accessed through any browser. An institution administrator with no blockchain knowledge can create and issue a blockchain-anchored certificate through a form interface in under five minutes.")

add_body(doc, "Blockcerts uses the Bitcoin blockchain for anchoring, which, while highly secure, has higher transaction fees and slower confirmation times than Polygon. Certichain Africa's use of Polygon makes high-volume certificate issuance economically feasible. Blockcerts does not include PDF generation, mobile money payment, or two-factor authentication.")

add_heading3(doc, "4.7.2 Comparison with Accredible and Credly")

add_body(doc, "Accredible and Credly are polished commercial platforms with strong track records in the European and North American markets. They offer well-designed credential displays, integration with LinkedIn and other professional networks, and reliable infrastructure. Certichain Africa is a student project and cannot match these platforms on production stability or feature breadth.")

add_body(doc, "However, Certichain Africa surpasses them on three dimensions that are critical for the African market. First, neither Accredible nor Credly offer blockchain anchoring: their credentials are stored on centralised servers and cannot be verified independently. Second, neither platform accepts mobile money payment: they require credit card or bank transfer, which excludes most Cameroonian institutions. Third, neither platform generates downloadable PDF certificates through a server-side pipeline: they display credentials as web pages or digital badges, which is not sufficient for institutions that need to provide physical-equivalent documentation.")

add_heading3(doc, "4.7.3 Summary Comparison Table")

tbl = doc.add_table(rows=7, cols=6)
tbl.style = 'Table Grid'
for i, h in enumerate(["Dimension", "Blockcerts", "Accredible", "Credly", "Sony/IBM", "Certichain Africa"]):
    tbl.rows[0].cells[i].text = h
    for p in tbl.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

comp_rows = [
    ("Blockchain Anchoring", "Yes", "No", "No", "Yes (private)", "Yes (Polygon)"),
    ("PDF Generation", "No", "Partial", "Partial", "No", "Yes (4 templates)"),
    ("MoMo Payment", "No", "No", "No", "No", "Yes (XAF)"),
    ("French Interface", "Partial", "No", "No", "No", "Yes"),
    ("Free Tier", "Self-host only", "No", "No", "No", "Yes (5/month)"),
    ("African Context Design", "No", "No", "No", "No", "Yes"),
]
for i, row_data in enumerate(comp_rows, start=1):
    for j, val in enumerate(row_data):
        tbl.rows[i].cells[j].text = val
        for p in tbl.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

add_fig(doc, "Table 4.11: Certichain Africa vs Existing Systems")

add_heading2(doc, "4.8 Evaluation Against Objectives")

add_body(doc, "This section evaluates the extent to which the implementation achieved each of the seven specific objectives stated in Chapter 1.")

tbl = doc.add_table(rows=8, cols=3)
tbl.style = 'Table Grid'
for i, h in enumerate(["Objective", "Implementation", "Status"]):
    tbl.rows[0].cells[i].text = h
    for p in tbl.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

obj_rows = [
    ("SO1: Secure institution authentication with OTP 2FA", "Implemented in app.py with OTP rate limiting, SHA-256 hashed OTP storage, 10-minute expiry, and 5-attempt lockout", "Achieved"),
    ("SO2: PDF generation for 3+ credential types with QR codes", "4 templates implemented in pdf_generator.py using ReportLab; QR codes embedded via qrcode[pil]", "Achieved"),
    ("SO3: Blockchain anchoring on Polygon network", "CertiChain.sol deployed; Web3.py integration in app.py; dual anchoring with SHA-256 fallback", "Achieved"),
    ("SO4: Pinata IPFS storage integration", "upload_to_ipfs() function implemented; local fallback when IPFS unavailable", "Achieved"),
    ("SO5: Public certificate verification system", "File upload, hash-based, and QR-code-linked verification implemented", "Achieved"),
    ("SO6: MTN MoMo subscription payment", "momo.py implements request_to_pay and get_status; plan upgrade on success", "Achieved"),
    ("SO7: French-capable interface and certificate templates", "All templates and routes use French text; PDF certificates output French prose", "Achieved"),
]

for i, (obj, impl, status) in enumerate(obj_rows, start=1):
    tbl.rows[i].cells[0].text = obj
    tbl.rows[i].cells[1].text = impl
    tbl.rows[i].cells[2].text = status
    for j in range(3):
        for p in tbl.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

add_fig(doc, "Table 4.10: Evaluation of Objectives Against Implementation")

add_body(doc, "All seven specific objectives were achieved in the implementation. Objectives SO1 through SO5 were fully achieved with the planned technical approach. SO6 was achieved in the sandbox environment, with production readiness dependent on completion of the MTN MoMo API production approval process. SO7 was achieved throughout the platform, though the degree of French-language completeness varies between the older template-based pages and the newer API-driven pages.")

add_heading2(doc, "4.9 Strengths and Limitations of the Implementation")

add_heading3(doc, "4.9.1 Strengths")

add_body(doc, "The dual anchoring strategy is one of the most significant architectural strengths of the implementation. By designing the system to always produce a verifiable certificate regardless of whether blockchain connectivity is available, the implementation avoids the single point of failure that plagues systems that depend entirely on external blockchain infrastructure. This resilience is especially valuable in the African deployment context where internet connectivity may be intermittent or where the issuer's Polygon wallet may temporarily run out of gas tokens.")

add_body(doc, "The four certificate templates represent a meaningful visual contribution. The diploma, certification, badge, and cyber diploma templates each have a distinct visual identity with a consistent brand language built around the navy, gold, and teal colour palette. The professional quality of these PDF outputs compares favourably with those produced by commercial platforms, and the ability to generate them entirely on the server side without any front-end design tools or external APIs is a practical advantage for a self-hosted deployment.")

add_body(doc, "The security design of the authentication system is another strength. The combination of password strength validation, PBKDF2-SHA256 hashing, OTP two-factor authentication, rate limiting, attempt counting, and session-based access control provides a level of security that is appropriate for a platform that controls the issuance of official academic credentials.")

add_body(doc, "The MTN Mobile Money integration is a genuine innovation in the context of African EdTech platforms. Very few open-source web applications include a working MoMo payment integration, and the implementation in momo.py provides a clean, reusable module that other developers in the region can study and adapt.")

add_heading3(doc, "4.9.2 Limitations")

add_body(doc, "The most significant limitation of the current implementation is the absence of a certificate revocation mechanism. Once a certificate is issued and anchored to the blockchain, there is no way to revoke it through the platform. In a real institutional deployment, situations arise where a certificate needs to be revoked: a student whose degree is rescinded, a fraudulent institution, or a data entry error. Addressing this limitation would require adding a revoke function to the smart contract and implementing a revocation check in the verification pipeline.")

add_body(doc, "The lack of institution KYC is a second important limitation. Any individual with an email address can create an institution account and begin issuing certificates. There is no mechanism to verify that the registering party is a legitimate educational institution. This means that a bad actor could use the platform to issue fake certificates from a fictitious institution. Addressing this would require a manual vetting process or integration with a government institution registry.")

add_body(doc, "The platform's dependence on Pinata for IPFS storage introduces a centralisation risk. While the blockchain record is truly decentralised, the certificate files themselves are stored on Pinata's servers. If Pinata changes its pricing model, goes offline, or removes files that are not being actively pinned, the certificates stored there may become inaccessible via IPFS. A more robust approach would pin certificate files on multiple IPFS nodes.")

add_body(doc, "The testing coverage is limited to manual functional testing. A production-grade platform should have an automated test suite with unit tests for each model method, integration tests for each API endpoint, and end-to-end tests for the complete certificate creation and verification workflows. The absence of automated tests means that regressions can be introduced by future code changes without being immediately detected.")

page_break(doc)

# ==============================================================================
# CHAPTER 5: CONCLUSION AND RECOMMENDATIONS
# ==============================================================================

add_heading1(doc, "CHAPTER 5\nCONCLUSION AND RECOMMENDATIONS")

add_heading2(doc, "5.1 Summary of Findings")

add_body(doc, "This project set out to design and implement a blockchain-based academic certificate management system adapted to the needs of educational institutions in Cameroon and the wider Francophone African region. The motivation for the project was clear: certificate fraud is a documented and serious problem in Africa, existing blockchain credential solutions are not designed for the African context, and no single platform combines the full set of features that African institutions need in order to adopt digital credentials meaningfully.")

add_body(doc, "The system that was designed and implemented, Certichain Africa, addresses all of these issues. It provides a complete certificate management workflow that allows an institution to register, log in securely with two-factor authentication, create professional PDF certificates for four different credential types, anchor those certificates to the Polygon blockchain or to a local SHA-256 fingerprint as a fallback, store the certificate files on IPFS, manage the certificates through a browser-based dashboard, allow any member of the public to verify a certificate by scanning a QR code or uploading the file, and pay for a subscription using MTN Mobile Money.")

add_body(doc, "The technical findings of the implementation process confirmed that each of these features is implementable using freely available, well-documented, open-source tools. Flask, Web3.py, ReportLab, Pinata's Python-compatible REST API, and the MTN MoMo Developer API are all accessible to a developer with standard Python web development skills and a modest budget. The total external service costs for operating the platform at low volume are near zero: the Pinata free tier covers the initial IPFS storage, the Polygon gas fees are negligible, and the MoMo sandbox is freely accessible for development and testing.")

add_body(doc, "The security findings were also positive. The authentication system, including the OTP mechanism, password hashing, rate limiting, and session management, withstood all the test scenarios that were designed to challenge it. The certificate verification system correctly identified genuine certificates and correctly rejected tampered ones. The blockchain anchoring mechanism produced verifiable on-chain records that can be independently queried by anyone with access to the Polygon network.")

add_heading2(doc, "5.2 Progress Made to Cover the Research Gap")

add_body(doc, "The research gap identified in Chapter 2 was stated as follows: no existing blockchain-based academic credential platform is designed specifically for the operational, linguistic, financial, and technical constraints of educational institutions in Francophone Sub-Saharan Africa. Certichain Africa addresses this gap on all four dimensions.")

add_body(doc, "On the financial dimension, the MTN Mobile Money payment integration and the free tier make the platform accessible to institutions that cannot pay with credit cards and cannot afford high subscription fees. On the linguistic dimension, the French-language interface and the French-text certificate templates make the platform genuinely usable by Francophone institutions. On the technical constraint dimension, the dual anchoring strategy and the local file storage fallback make the platform functional in environments where internet connectivity is intermittent. On the completeness dimension, the single platform provides everything an institution needs from registration to certificate issuance to public verification, without requiring the use of multiple tools or custom development.")

add_body(doc, "The gap is not fully closed by a single student project. Certichain Africa is a proof of concept that demonstrates what is possible; it is not a production-deployed, internationally recognised credential standard. The gap will be more fully addressed as the platform matures through continued development, institutional adoption, and regulatory engagement.")

add_heading2(doc, "5.3 Conclusions")

add_body(doc, "This project reaches three main conclusions.")

add_body(doc, "The first conclusion is that blockchain-based academic credential management is technically feasible for African educational institutions using accessible, low-cost tools. The implementation demonstrated that a developer with standard Python web skills can build a working blockchain credential system without specialised blockchain expertise, using Web3.py to interact with the Polygon network and Pinata to handle IPFS storage. The technical barrier to entry for blockchain credential management is lower than it is commonly perceived to be.")

add_body(doc, "The second conclusion is that the African context requires specific design decisions that are not addressed by existing international solutions. Mobile money payment, French-language support, connectivity resilience, and a complete out-of-the-box workflow are not nice-to-have features in this context; they are prerequisites for adoption. A platform that is designed from the beginning with these constraints in mind, as Certichain Africa was, is far more likely to be adopted by African institutions than one that requires adaptation from a different context.")

add_body(doc, "The third conclusion is that Design Science Research is an appropriate methodology for this type of project. The artefact produced by the project, the Certichain Africa platform, is a genuine contribution to both the technical and the social problem it addresses. The evaluation of the artefact against the defined requirements provides empirical evidence of its effectiveness, and the documentation of the design decisions provides knowledge that goes beyond the specific artefact itself.")

add_heading2(doc, "5.4 Recommendations and Future Work")

add_body(doc, "Based on the findings, limitations, and conclusions of this project, the following recommendations are made for future development of Certichain Africa and for related work in this area.")

add_heading3(doc, "5.4.1 Certificate Revocation")

add_body(doc, "The highest priority for future development is the implementation of a certificate revocation mechanism. This would require adding a revokeCertificate function to the CertiChain smart contract that sets a revoked flag on the certificate record, and updating the verification pipeline to check the revoked flag and return an appropriate response when a revoked certificate is queried. Revocation events should also be recorded in the platform's database so that institutions can maintain an audit trail of which certificates have been revoked and why.")

add_heading3(doc, "5.4.2 Institution KYC and Verification")

add_body(doc, "A manual or automated institution verification process should be implemented before the platform is deployed in a production environment. At minimum, this should require institutions to upload a registration document or official letter proving their status as a legal educational entity. More advanced approaches could include integration with national education ministry registries or with identity verification services such as Smile Identity, which specialises in African digital identity.")

add_heading3(doc, "5.4.3 Student-Facing Portal")

add_body(doc, "A student-facing portal would allow certificate recipients to access their credentials directly without needing to contact the issuing institution. Students could log in, view all the certificates that have been issued to them across multiple institutions, share individual certificates via a public link, and download the PDFs. This would significantly enhance the platform's value for certificate recipients and would align Certichain Africa more closely with commercial platforms like Accredible and Credly on the recipient experience dimension.")

add_heading3(doc, "5.4.4 Mobile Application")

add_body(doc, "A mobile application for Android and iOS would increase the accessibility of the platform for institution administrators who may work primarily from smartphones rather than desktop computers. The application should support certificate creation, management, and QR-code-based verification. Given that many African educators rely on mobile internet rather than fixed broadband, a mobile-first application could significantly increase adoption.")

add_heading3(doc, "5.4.5 Automated Test Suite")

add_body(doc, "A comprehensive automated test suite should be developed using Python's pytest framework. The suite should include unit tests for all model methods, integration tests for all API endpoints, and end-to-end tests for the complete certificate creation and verification workflows. Automated testing would reduce the risk of regressions and make it safer for multiple developers to contribute to the codebase.")

add_heading3(doc, "5.4.6 Additional Payment Methods")

add_body(doc, "Orange Money, which is MTN MoMo's main competitor in Central and West Africa, should be integrated as an additional payment method. Card payment support through a provider such as Flutterwave or Paystack, which specialise in African markets, would further broaden accessibility. Wave Mobile Money, which is dominant in Senegal and other West African markets, would enable expansion beyond the initial Cameroon and Central Africa focus.")

add_heading3(doc, "5.4.7 Bulk Certificate Issuance")

add_body(doc, "The ability to issue certificates in bulk, for example by uploading a CSV file containing the details of a graduating class, would significantly reduce the time and effort required for large-scale certificate issuance. This feature is particularly important for universities that award hundreds or thousands of degrees at each graduation ceremony.")

add_heading2(doc, "5.5 Final Conclusion")

add_body(doc, "This project has successfully designed and implemented a blockchain-based academic certificate management system that is adapted to the needs of educational institutions in Cameroon and the broader Francophone African region. The Certichain Africa platform demonstrates that the combination of the Flask web framework, the Polygon blockchain, the Pinata IPFS service, the ReportLab PDF generation library, and the MTN Mobile Money API can produce a coherent, functional, and locally relevant certificate management system that addresses the identified research gap.")

add_body(doc, "The certificate fraud crisis in Africa is real, and the consequences for honest graduates, responsible institutions, and employers who rely on the integrity of academic credentials are significant. Technology alone cannot solve this problem, but it can provide tools that make fraud much harder and verification much easier. Certichain Africa is one such tool, and this report provides a complete account of how it was designed, why those design decisions were made, and what the system achieves.")

add_body(doc, "It is hoped that this work will serve as both a practical reference for developers and institutions in the region who wish to implement similar systems, and as a contribution to the academic conversation about the role of blockchain technology in addressing educational challenges in Sub-Saharan Africa. The source code for Certichain Africa is available in the project repository and is designed to be understood, extended, and deployed by any developer with standard Python web development skills.")

add_body(doc, "The journey from identifying a problem to delivering a working solution that addresses it is the essence of engineering. This project has made that journey, and the result is a platform that has the potential, with continued development and institutional adoption, to make a meaningful contribution to the integrity of academic credentials in Africa.")

page_break(doc)

# ==============================================================================
# APPENDICES
# ==============================================================================

add_heading1(doc, "APPENDICES")

add_heading2(doc, "Appendix A: Key Code Listings")

add_heading3(doc, "A.1 Database Models (models.py)")

add_code(doc, "models.py", """from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime

db = SQLAlchemy()

class Institution(db.Model):
    __tablename__ = 'institutions'
    id             = db.Column(db.Integer, primary_key=True)
    name           = db.Column(db.String(255), nullable=False)
    email          = db.Column(db.String(255), unique=True, nullable=False, index=True)
    password_hash  = db.Column(db.String(255), nullable=False)
    is_verified    = db.Column(db.Boolean, default=True)
    created_at     = db.Column(db.DateTime, default=datetime.utcnow)
    wallet_address = db.Column(db.String(255))
    plan           = db.Column(db.String(50), default='free')
    plan_expires_at= db.Column(db.DateTime, nullable=True)
    certificates   = db.relationship('Certificate', back_populates='institution',
                                     cascade='all, delete-orphan')
    payments       = db.relationship('Payment', back_populates='institution',
                                     cascade='all, delete-orphan')

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def cert_quota(self):
        limits = {'free': 5, 'starter': 50, 'academique': 300, 'enterprise': 999999}
        limit  = limits.get(self.plan or 'free', 5)
        now    = datetime.utcnow()
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        used = Certificate.query.filter(
            Certificate.institution_id == self.id,
            Certificate.created_at >= month_start,
        ).count()
        return {'used': used, 'limit': limit}""")

add_heading3(doc, "A.2 Smart Contract (CertiChain.sol)")

add_code(doc, "CertiChain.sol", """// SPDX-License-Identifier: MIT
pragma solidity ^0.8.0;

contract CertiChain {
    struct Certificate {
        string   ipfsHash;
        string   recipientName;
        uint256  issueDate;
        bool     exists;
    }

    mapping(bytes32 => Certificate) public certificates;
    address public owner;

    event CertificateIssued(
        bytes32 indexed certificateId,
        string  recipientName,
        uint256 issueDate
    );

    constructor() { owner = msg.sender; }

    modifier onlyOwner() {
        require(msg.sender == owner, "Not authorized");
        _;
    }

    function issueCertificate(
        bytes32 _certificateId,
        string memory _ipfsHash,
        string memory _recipientName
    ) public onlyOwner {
        require(!certificates[_certificateId].exists, "Already exists");
        certificates[_certificateId] = Certificate(
            _ipfsHash, _recipientName, block.timestamp, true
        );
        emit CertificateIssued(_certificateId, _recipientName, block.timestamp);
    }

    function verifyCertificate(bytes32 _certificateId)
        public view returns (bool) {
        return certificates[_certificateId].exists;
    }

    function getCertificate(bytes32 _certificateId)
        public view
        returns (string memory ipfsHash, string memory recipientName, uint256 issueDate)
    {
        Certificate memory cert = certificates[_certificateId];
        return (cert.ipfsHash, cert.recipientName, cert.issueDate);
    }
}""")

add_heading3(doc, "A.3 MTN MoMo Client (momo.py)")

add_code(doc, "momo.py", """import os, base64, requests
from dotenv import load_dotenv
load_dotenv()

BASE_URL    = os.getenv('MOMO_BASE_URL', 'https://sandbox.momodeveloper.mtn.com')
SUB_KEY     = os.getenv('MOMO_SUBSCRIPTION_KEY', '')
API_USER    = os.getenv('MOMO_API_USER', '')
API_KEY     = os.getenv('MOMO_API_KEY', '')
ENVIRONMENT = os.getenv('MOMO_ENVIRONMENT', 'sandbox')

def _token():
    credentials = base64.b64encode(f"{API_USER}:{API_KEY}".encode()).decode()
    r = requests.post(f"{BASE_URL}/collection/token/",
        headers={"Authorization": f"Basic {credentials}",
                 "Ocp-Apim-Subscription-Key": SUB_KEY}, timeout=15)
    r.raise_for_status()
    return r.json()['access_token']

def request_to_pay(phone, amount, currency, reference_id, note):
    token = _token()
    r = requests.post(
        f"{BASE_URL}/collection/v1_0/requesttopay",
        json={"amount": str(amount), "currency": currency,
              "externalId": reference_id,
              "payer": {"partyIdType": "MSISDN", "partyId": phone},
              "payerMessage": note, "payeeNote": note},
        headers={"Authorization": f"Bearer {token}",
                 "X-Target-Environment": ENVIRONMENT,
                 "Ocp-Apim-Subscription-Key": SUB_KEY,
                 "X-Reference-Id": reference_id,
                 "Content-Type": "application/json"},
        timeout=20)
    return r.status_code == 202""")

add_heading2(doc, "Appendix B: Diagram Descriptions")

add_body(doc, "All diagrams referenced in this report should be drawn using a diagramming tool such as draw.io (app.diagrams.net) or Lucidchart. Detailed descriptions of each diagram are provided in the relevant sections of Chapters 3 and 4. The descriptions specify every actor, entity, attribute, relationship, arrow direction, and label needed to produce each diagram. Once drawn, the diagram images should be inserted into the report at the marked [FIGURE] placeholders.")

add_heading2(doc, "Appendix C: Installation and Setup Guide")

add_body(doc, "The following instructions describe how to set up and run Certichain Africa on a local development machine. These instructions assume a Windows environment with Python 3.8 or later installed.")

add_heading3(doc, "C.1 Clone and Set Up the Environment")

add_code(doc, "PowerShell", """# Navigate to the project directory
cd C:\\path\\to\\Certichain-Africa

# Create a virtual environment
python -m venv venv

# Activate the virtual environment
.\\venv\\Scripts\\Activate.ps1

# Install dependencies
pip install -r requirements.txt""")

add_heading3(doc, "C.2 Configure Environment Variables")

add_code(doc, ".env (create in project root)", """# Database
DATABASE_URL=sqlite:///certichain.db

# Email (SMTP)
MAIL_SERVER=smtp.gmail.com
MAIL_PORT=587
MAIL_USE_TLS=True
MAIL_USERNAME=your_email@gmail.com
MAIL_PASSWORD=your_app_password
MAIL_DEFAULT_SENDER=your_email@gmail.com

# Blockchain (Polygon)
POLYGON_RPC=https://polygon-rpc.com
CONTRACT_ADDRESS=0xYourContractAddress
ISSUER_ADDRESS=0xYourWalletAddress
ISSUER_PRIVATE_KEY=YourPrivateKey

# IPFS (Pinata)
PINATA_API_KEY=YourPinataApiKey
PINATA_SECRET_KEY=YourPinataSecretKey

# MTN MoMo
MOMO_SUBSCRIPTION_KEY=YourSubscriptionKey
MOMO_API_USER=YourApiUser
MOMO_API_KEY=YourApiKey
MOMO_ENVIRONMENT=sandbox""")

add_heading3(doc, "C.3 Run the Application")

add_code(doc, "PowerShell", """# Run the Flask application
python app.py

# The application will be available at:
# http://localhost:5000""")

page_break(doc)

# ==============================================================================
# REFERENCES
# ==============================================================================

add_heading1(doc, "REFERENCES")

references = [
    "[1]  African Union Commission, \"Education, Science and Technology in Africa: Continental Education Strategy,\" African Union, Addis Ababa, 2019.",
    "[2]  S. Nakamoto, \"Bitcoin: A Peer-to-Peer Electronic Cash System,\" 2008. [Online]. Available: https://bitcoin.org/bitcoin.pdf",
    "[3]  Learning Machine Technologies, \"Blockcerts: An Open Infrastructure for Academic Credentials on the Blockchain,\" MIT Media Lab, Cambridge, MA, 2016.",
    "[4]  GSMA, \"State of the Industry Report on Mobile Money 2022,\" GSMA, London, 2022.",
    "[5]  P. Sharma and A. Bhardwaj, \"Digital Credentials and Blockchain: A Review of Current Practices,\" Journal of Information Technology Education, vol. 19, pp. 211-235, 2020.",
    "[6]  IMS Global Learning Consortium, \"Open Badges 2.0 Specification,\" IMS Global, Lake Mary, FL, 2018. [Online]. Available: https://www.imsglobal.org/sites/default/files/Badges/OBv2p0Final/index.html",
    "[7]  V. Buterin, \"Ethereum: A Next-Generation Smart Contract and Decentralized Application Platform,\" Ethereum Foundation, 2014. [Online]. Available: https://ethereum.org/en/whitepaper/",
    "[8]  Polygon Technology, \"Polygon PoS: A Proof of Stake Sidechain for Ethereum,\" Polygon, 2021. [Online]. Available: https://docs.polygon.technology/",
    "[9]  G. Wood, \"Ethereum: A Secure Decentralised Generalised Transaction Ledger (Yellow Paper),\" Ethereum Foundation, 2014.",
    "[10] J. Benet, \"IPFS - Content Addressed, Versioned, P2P File System,\" arXiv preprint arXiv:1407.3561, 2014.",
    "[11] Pinata Technologies, \"Pinata IPFS Pinning Service Documentation,\" 2023. [Online]. Available: https://docs.pinata.cloud/",
    "[12] National Institute of Standards and Technology, \"Secure Hash Standard (SHS),\" FIPS PUB 180-4, NIST, Gaithersburg, MD, 2015.",
    "[13] OWASP Foundation, \"OWASP Top Ten 2021: The Ten Most Critical Web Application Security Risks,\" OWASP, 2021. [Online]. Available: https://owasp.org/Top10/",
    "[14] A. R. Hevner, S. T. March, J. Park, and S. Ram, \"Design Science in Information Systems Research,\" MIS Quarterly, vol. 28, no. 1, pp. 75-105, 2004.",
    "[15] K. Peffers, T. Tuunanen, M. A. Rothenberger, and S. Chatterjee, \"A Design Science Research Methodology for Information Systems Research,\" Journal of Management Information Systems, vol. 24, no. 3, pp. 45-77, 2007.",
    "[16] F. D. Davis, \"Perceived Usefulness, Perceived Ease of Use, and User Acceptance of Information Technology,\" MIS Quarterly, vol. 13, no. 3, pp. 319-340, 1989.",
    "[17] E. M. Rogers, Diffusion of Innovations, 5th ed. New York, NY: Free Press, 2003.",
    "[18] A. Grech and A. F. Camilleri, \"Blockchain in Education,\" European Commission, Publications Office of the European Union, Luxembourg, 2017.",
    "[19] M. Turkanovic, M. Holbl, K. Kosic, M. Hericko, and A. Kamisalic, \"EduCTX: A Blockchain-Based Higher Education Credit Platform,\" IEEE Access, vol. 6, pp. 5112-5127, 2018.",
    "[20] D. Lizcano, J. A. Lara, B. White, and S. Aljawarneh, \"Blockchain-Based Approach to Create a Model of Trust in Open and Decentralized Educational Environments,\" Journal of Computing in Higher Education, vol. 32, no. 1, pp. 109-134, 2020.",
    "[21] O. S. Adesina, \"Fake Certificates and the Challenge of Professional Conduct in Nigeria,\" Journal of Education and Practice, vol. 2, no. 11-12, pp. 43-52, 2011.",
    "[22] K. Adekunle, A. M. Ojo, and B. Adewale, \"Certificate Fraud and Human Resource Integrity in Ghanaian Organizations,\" African Journal of Business and Economic Research, vol. 13, no. 2, pp. 67-84, 2018.",
    "[23] W. Jack and T. Suri, \"Mobile Money: The Economics of M-Pesa,\" National Bureau of Economic Research, Working Paper 16721, Cambridge, MA, 2011.",
    "[24] Adobe Systems Incorporated, \"PDF Reference, Sixth Edition: Adobe Portable Document Format Version 1.7,\" Adobe Systems, San Jose, CA, 2006.",
    "[25] ReportLab Inc., \"ReportLab Open Source PDF Library Documentation,\" 2023. [Online]. Available: https://www.reportlab.com/docs/reportlab-userguide.pdf",
    "[26] Accredible, \"Accredible Digital Certificate and Badge Platform,\" 2023. [Online]. Available: https://www.accredible.com",
    "[27] Credly Inc., \"Credly: The World's Leading Digital Credentialing Platform,\" 2023. [Online]. Available: https://info.credly.com",
    "[28] Sony Global Education, \"Sony Global Education Develops Technology Using Blockchain,\" Press Release, Sony Corporation, Tokyo, 2017.",
    "[29] K. Beck et al., \"Manifesto for Agile Software Development,\" 2001. [Online]. Available: https://agilemanifesto.org",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.line_spacing  = Pt(20)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ── Save final merged document ────────────────────────────────────────────────
doc.save('docs/FYP_REPORT.docx')
print("Final report saved to docs/FYP_REPORT.docx")
