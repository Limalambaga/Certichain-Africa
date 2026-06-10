"""
Certichain Africa - FYP Report Generator Part 1
Covers: Front Matter, Chapter 1, Chapter 2, Chapter 3
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Default body font ─────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(6)

def set_font(run, bold=False, italic=False, size=12, color=None):
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_body(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(text.upper())
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(14)
    run.font.bold  = True
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(12)
    run.font.bold  = True
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(12)
    run.font.bold   = True
    run.font.italic = True
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(1.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_figure_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(f'[{text}]')
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    return p

def page_break(doc):
    doc.add_page_break()

# ==============================================================================
# TITLE PAGE
# ==============================================================================

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("ICT UNIVERSITY")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Faculty of Engineering and Technology")
run.font.name = 'Times New Roman'
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Department of Information Systems and Networking")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FINAL YEAR PROJECT REPORT")
run.font.name = 'Times New Roman'
run.font.size = Pt(13)
run.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A Blockchain-Based Academic Certificate Issuance\nand Verification System")
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted by:")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LIMALA MBAGA PAUL BETSALEEL")
run.font.name = 'Times New Roman'
run.font.size = Pt(13)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Student ID: ICTU20223224")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Supervised by:")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dr. NKANDEU PASCAL")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Co-Supervisor(s): [TO BE FILLED]")
run.font.name   = 'Times New Roman'
run.font.size   = Pt(12)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("In partial fulfillment of the requirements for the award of the\nBachelor of Science Degree in Information Systems and Networking")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Academic Year: 2025 - 2026")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True

page_break(doc)

# ==============================================================================
# DECLARATION
# ==============================================================================

add_heading1(doc, "DECLARATION OF ORIGINALITY")

add_body(doc, "I, LIMALA MBAGA PAUL BETSALEEL, student of ICT University with Student ID ICTU20223224, hereby declare that this Final Year Project report titled \"A Blockchain-Based Academic Certificate Issuance and Verification System\" is my own original work. It has been carried out under the supervision of Dr. NKANDEU PASCAL in the Department of Information Systems and Networking.")

add_body(doc, "I further declare that this work has not been previously submitted in whole or in part for the award of any degree or other qualification at this university or any other institution. All sources of information used in this project have been duly acknowledged and referenced in accordance with the academic conventions of ICT University.")

add_body(doc, "Where the work of other authors has been used to support arguments or to provide background information, such work has been explicitly cited and listed in the bibliography. Any resemblance to previously published works is purely coincidental and unintentional.")

add_body(doc, "I understand that any act of plagiarism or academic dishonesty may result in the cancellation of this work and the imposition of appropriate disciplinary sanctions as provided for in the university's academic integrity policy.")

doc.add_paragraph()
add_body(doc, "Signature: ___________________________          Date: ___________________")
doc.add_paragraph()
add_body(doc, "LIMALA MBAGA PAUL BETSALEEL")
add_body(doc, "ICTU20223224")

page_break(doc)

# ==============================================================================
# CERTIFICATION
# ==============================================================================

add_heading1(doc, "CERTIFICATION BY SUPERVISOR")

add_body(doc, "This is to certify that this Final Year Project report titled \"A Blockchain-Based Academic Certificate Issuance and Verification System,\" submitted by LIMALA MBAGA PAUL BETSALEEL (Student ID: ICTU20223224), has been prepared under my direct supervision and guidance in the Department of Information Systems and Networking at ICT University, Cameroon.")

add_body(doc, "To the best of my knowledge, the work presented in this report is original and has not been submitted elsewhere for the award of any degree or academic qualification. The report is ready for examination by the Faculty.")

doc.add_paragraph()
add_body(doc, "Supervisor:")
add_body(doc, "Name: Dr. NKANDEU PASCAL")
add_body(doc, "Title: [SUPERVISOR TITLE]")
add_body(doc, "Department: Information Systems and Networking")
add_body(doc, "ICT University, Cameroon")
doc.add_paragraph()
add_body(doc, "Signature: ___________________________          Date: ___________________")

doc.add_paragraph()
doc.add_paragraph()

add_body(doc, "Co-Supervisor (if applicable):")
add_body(doc, "Name: [CO-SUPERVISOR NAME]")
add_body(doc, "Title: [CO-SUPERVISOR TITLE]")
add_body(doc, "Signature: ___________________________          Date: ___________________")

page_break(doc)

# ==============================================================================
# FACULTY APPROVAL
# ==============================================================================

add_heading1(doc, "FACULTY APPROVAL PAGE")

add_body(doc, "This Final Year Project report titled \"A Blockchain-Based Academic Certificate Issuance and Verification System\" submitted by LIMALA MBAGA PAUL BETSALEEL (Student ID: ICTU20223224) has been examined and approved by the Department of Information Systems and Networking, Faculty of Engineering and Technology, ICT University, Cameroon.")

doc.add_paragraph()

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
headers = ["Role", "Name", "Signature and Date"]
for i, h in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

rows_data = [
    ("Head of Department", "[HOD NAME]", "________________"),
    ("Internal Examiner", "[EXAMINER NAME]", "________________"),
    ("External Examiner", "[EXAMINER NAME]", "________________"),
]
for i, (role, name, sig) in enumerate(rows_data, start=1):
    tbl.rows[i].cells[0].text = role
    tbl.rows[i].cells[1].text = name
    tbl.rows[i].cells[2].text = sig
    for c in tbl.rows[i].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

page_break(doc)

# ==============================================================================
# ACKNOWLEDGMENTS
# ==============================================================================

add_heading1(doc, "ACKNOWLEDGMENTS")

add_body(doc, "First and foremost, I give thanks to the Almighty God for the strength, wisdom, and perseverance He granted me throughout this entire academic journey. Without His grace, none of this would have been possible.")

add_body(doc, "I owe my deepest gratitude to my supervisor, Dr. NKANDEU PASCAL, whose guidance, patience, and intellectual generosity shaped every chapter of this work. His ability to challenge my thinking while remaining encouraging made the difference between a shallow project and one I am genuinely proud of. I am truly grateful for the time and expertise he devoted to this research.")

add_body(doc, "I would also like to express my sincere thanks to the entire faculty of the Department of Information Systems and Networking at ICT University. The knowledge I received in the lecture halls, laboratories, and informal conversations with lecturers provided the foundation on which this project was built. Their commitment to excellence is reflected in every page of this report.")

add_body(doc, "To my classmates and study partners, especially those who endured late-night debugging sessions and helped me think through difficult design decisions, I say a heartfelt thank you. The collaborative spirit we maintained throughout this programme is something I will carry forward into my professional life.")

add_body(doc, "I am deeply grateful to my family for their unconditional love and support. To my parents, whose sacrifices made my education possible, and to my siblings who kept me grounded and motivated during moments of doubt, I dedicate this work to you.")

add_body(doc, "Finally, I acknowledge the broader open-source communities behind Flask, Web3.py, ReportLab, Pinata, and the Polygon Network. The tools they have built and freely shared made it possible for a student with limited resources to design and implement a system of this ambition.")

page_break(doc)

# ==============================================================================
# ABSTRACT
# ==============================================================================

add_heading1(doc, "ABSTRACT")

add_body(doc, "The integrity of academic credentials is a critical concern for educational institutions, employers, and governments across Africa. Certificate fraud, including the falsification and forgery of academic documents, has undermined trust in educational systems and created serious consequences for both individuals and organisations that rely on those credentials. Existing solutions are largely paper-based, manual, or dependent on centralised verification authorities that are not always accessible or trustworthy in the African context.")

add_body(doc, "This project presents the design, development, and implementation of Certichain Africa, a web-based platform that enables educational institutions to issue, manage, and verify academic certificates using blockchain technology and distributed file storage. The system is built on the Python Flask framework and uses the Polygon blockchain network together with the Pinata IPFS service to anchor and store certificates in a manner that is tamper-proof and publicly verifiable. Certificates are generated as professionally designed PDF files using the ReportLab library, and each PDF embeds a QR code that links to a public verification page. The platform supports three categories of credentials, namely diplomas, professional certifications, and skill badges, along with a specialised cybersecurity diploma template. A two-factor authentication mechanism based on one-time passwords delivered by email ensures that only authorised institutions can access the system. Subscription payments are processed through the MTN Mobile Money API, making the platform accessible in the Central African economic context where card payments are not widely available.")

add_body(doc, "The system was developed following the Design Science Research methodology combined with Agile principles, and was evaluated against the functional and non-functional requirements identified during the analysis phase. Results show that the platform successfully anchors certificates to the blockchain, verifies their authenticity, and generates downloadable PDF credentials in under two seconds per certificate. Certichain Africa addresses a clear gap in the African EdTech landscape by providing a locally adapted, cost-effective, and technically sound solution to the certificate fraud problem.")

doc.add_paragraph()
add_body(doc, "Keywords: blockchain, academic certificate, IPFS, Polygon, Flask, digital credentials, certificate fraud, Mobile Money, Africa, EdTech", italic=True)

page_break(doc)

# RESUME (French)
add_heading1(doc, "RESUME")

add_body(doc, "L'integrite des diplomes et certificats academiques est une preoccupation majeure pour les etablissements d'enseignement, les employeurs et les gouvernements a travers l'Afrique. La fraude aux diplomes, qui comprend la falsification et la contrefacon de documents academiques, a erode la confiance dans les systemes educatifs et engendre des consequences graves pour les individus et les organisations qui s'appuient sur ces documents.")

add_body(doc, "Ce projet presente la conception, le developpement et la mise en oeuvre de Certichain Africa, une plateforme web qui permet aux etablissements d'enseignement d'emettre, de gerer et de verifier des certificats academiques a l'aide de la technologie blockchain et du stockage de fichiers distribue. Le systeme est developpe avec le framework Python Flask et utilise le reseau blockchain Polygon ainsi que le service Pinata IPFS pour ancrer et stocker les certificats de maniere infalsifiable et verifiable publiquement. Les certificats sont generes sous forme de fichiers PDF professionnels a l'aide de la bibliotheque ReportLab, et chaque PDF integre un code QR renvoyant vers une page de verification publique. La plateforme prend en charge trois categories de diplomes, a savoir les diplomes, les certifications professionnelles et les badges de competences, ainsi qu'un modele de diplome specialise en cybersecurite. Un mecanisme d'authentification a deux facteurs base sur des mots de passe a usage unique envoyes par email garantit que seules les institutions autorisees peuvent acceder au systeme. Les paiements d'abonnement sont traites via l'API MTN Mobile Money, rendant la plateforme accessible dans le contexte economique de l'Afrique centrale.")

add_body(doc, "Le systeme a ete developpe selon la methodologie de la recherche en conception (Design Science Research) combinee aux principes Agile, et a ete evalue par rapport aux exigences fonctionnelles et non fonctionnelles identifiees lors de la phase d'analyse. Certichain Africa comble une lacune evidente dans le paysage EdTech africain en proposant une solution localement adaptee, rentable et techniquement rigoureuse au probleme de la fraude aux diplomes.")

doc.add_paragraph()
add_body(doc, "Mots-cles : blockchain, certificat academique, IPFS, Polygon, Flask, diplome numerique, fraude aux diplomes, Mobile Money, Afrique, EdTech", italic=True)

page_break(doc)

# ==============================================================================
# TABLE OF CONTENTS (manual)
# ==============================================================================

add_heading1(doc, "TABLE OF CONTENTS")

toc_entries = [
    ("Declaration of Originality", "ii"),
    ("Certification by Supervisor", "iii"),
    ("Faculty Approval Page", "iv"),
    ("Acknowledgments", "v"),
    ("Abstract", "vi"),
    ("Resume", "vii"),
    ("Table of Contents", "viii"),
    ("List of Figures", "x"),
    ("List of Tables", "xi"),
    ("List of Abbreviations", "xii"),
    ("CHAPTER 1: INTRODUCTION", "1"),
    ("    1.1  Background and Context", "1"),
    ("    1.2  Problem Statement", "4"),
    ("    1.3  Objectives of the Study", "6"),
    ("    1.4  Research Questions", "7"),
    ("    1.5  Scope and Limitations", "8"),
    ("    1.6  Significance of the Study", "9"),
    ("    1.7  Organization of the Report", "11"),
    ("CHAPTER 2: LITERATURE REVIEW", "12"),
    ("    2.1  Introduction", "12"),
    ("    2.2  Conceptual Literature Review", "12"),
    ("    2.3  Theoretical Framework", "18"),
    ("    2.4  Empirical Literature Review", "21"),
    ("    2.5  Case Studies of Existing Platforms", "25"),
    ("    2.6  Research Gap", "28"),
    ("    2.7  Conclusion", "30"),
    ("CHAPTER 3: METHODOLOGY", "31"),
    ("    3.1  Introduction", "31"),
    ("    3.2  Research Design", "31"),
    ("    3.3  Tools and Technologies Used", "34"),
    ("    3.4  System Design Approach", "40"),
    ("    3.5  Data Collection and Requirements Gathering", "42"),
    ("    3.6  Software Development Life Cycle Phases", "44"),
    ("    3.7  Data Flow Diagram", "47"),
    ("    3.8  Conclusion", "49"),
    ("CHAPTER 4 (PART 1): IMPLEMENTATION AND RESULTS", "50"),
    ("    4.1  System Analysis", "50"),
    ("    4.2  System Design", "53"),
    ("    4.3  Implementation", "62"),
    ("    4.4  Quality Assurance and Testing", "88"),
    ("    4.5  Application Overview", "92"),
    ("CHAPTER 4 (PART 2): DISCUSSION", "95"),
    ("    4.6  Interpretation of Results", "95"),
    ("    4.7  Comparison with Existing Works", "97"),
    ("    4.8  Evaluation Against Objectives", "100"),
    ("    4.9  Strengths and Limitations", "102"),
    ("CHAPTER 5: CONCLUSION AND RECOMMENDATIONS", "105"),
    ("    5.1  Summary of Findings", "105"),
    ("    5.2  Progress Made to Cover the Research Gap", "107"),
    ("    5.3  Conclusions", "108"),
    ("    5.4  Recommendations and Future Work", "109"),
    ("    5.5  Final Conclusion", "111"),
    ("APPENDICES", "112"),
    ("    Appendix A: Key Code Listings", "112"),
    ("    Appendix B: Diagram Descriptions", "120"),
    ("    Appendix C: Installation Guide", "122"),
    ("REFERENCES", "124"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(entry)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = entry.startswith("CHAPTER") or entry.startswith("APPENDICES") or entry.startswith("REFERENCES")
    tab_run = p.add_run(f"\t{page}")
    tab_run.font.name = 'Times New Roman'
    tab_run.font.size = Pt(11)

page_break(doc)

# ==============================================================================
# LIST OF FIGURES
# ==============================================================================

add_heading1(doc, "LIST OF FIGURES")

figures = [
    ("Figure 3.1", "Design Science Research Process Applied to Certichain Africa", "33"),
    ("Figure 3.2", "Agile Sprint Breakdown for Certichain Africa Development", "35"),
    ("Figure 3.3", "Level-0 Data Flow Diagram (Context Diagram)", "48"),
    ("Figure 3.4", "Level-1 Data Flow Diagram", "49"),
    ("Figure 4.1", "Use Case Diagram for Certichain Africa", "55"),
    ("Figure 4.2", "Class Diagram for Certichain Africa", "58"),
    ("Figure 4.3", "Entity-Relationship Diagram", "60"),
    ("Figure 4.4", "System Architecture Diagram (MVC + External Services)", "62"),
    ("Figure 4.5", "Certificate Issuance Sequence Diagram", "70"),
    ("Figure 4.6", "Authentication Flow Sequence Diagram", "65"),
    ("Figure 4.7", "Landing Page Screenshot", "92"),
    ("Figure 4.8", "Login Page Screenshot", "93"),
    ("Figure 4.9", "OTP Verification Page Screenshot", "93"),
    ("Figure 4.10", "Institution Dashboard Screenshot", "93"),
    ("Figure 4.11", "Certificate Creation Form Screenshot", "94"),
    ("Figure 4.12", "Generated Diploma PDF Sample", "94"),
    ("Figure 4.13", "Certificate Management Page Screenshot", "94"),
    ("Figure 4.14", "Public Verification Page Screenshot", "95"),
    ("Figure 4.15", "Pricing and Subscription Page Screenshot", "95"),
    ("Figure 4.16", "Settings Page Screenshot", "95"),
]

for fig, caption, page in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(f"{fig}:  {caption}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    tab = p.add_run(f"\t{page}")
    tab.font.name = 'Times New Roman'
    tab.font.size = Pt(11)

page_break(doc)

# ==============================================================================
# LIST OF TABLES
# ==============================================================================

add_heading1(doc, "LIST OF TABLES")

tables = [
    ("Table 2.1", "Comparison of Existing Certificate Verification Platforms", "26"),
    ("Table 3.1", "Tools and Technologies Used in Certichain Africa", "34"),
    ("Table 3.2", "Agile Sprint Plan and Deliverables", "45"),
    ("Table 4.1", "Functional Requirements of Certichain Africa", "51"),
    ("Table 4.2", "Non-Functional Requirements of Certichain Africa", "52"),
    ("Table 4.3", "Use Case Descriptions", "56"),
    ("Table 4.4", "Institution Entity Attributes", "59"),
    ("Table 4.5", "Certificate Entity Attributes", "59"),
    ("Table 4.6", "Payment Entity Attributes", "60"),
    ("Table 4.7", "OTP Rate-Limiting Parameters", "66"),
    ("Table 4.8", "Subscription Plan Features and Pricing (XAF)", "83"),
    ("Table 4.9", "Test Cases and Results", "89"),
    ("Table 4.10", "Evaluation of Objectives Against Implementation", "101"),
    ("Table 4.11", "Certichain Africa vs Existing Systems", "98"),
]

for tbl_id, caption, page in tables:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(f"{tbl_id}:  {caption}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    tab = p.add_run(f"\t{page}")
    tab.font.name = 'Times New Roman'
    tab.font.size = Pt(11)

page_break(doc)

# ==============================================================================
# LIST OF ABBREVIATIONS
# ==============================================================================

add_heading1(doc, "LIST OF ABBREVIATIONS")

abbrevs = [
    ("ABI",   "Application Binary Interface"),
    ("API",   "Application Programming Interface"),
    ("CORS",  "Cross-Origin Resource Sharing"),
    ("CRUD",  "Create, Read, Update, Delete"),
    ("DApp",  "Decentralised Application"),
    ("DB",    "Database"),
    ("DFD",   "Data Flow Diagram"),
    ("DSR",   "Design Science Research"),
    ("EdTech","Education Technology"),
    ("ER",    "Entity-Relationship"),
    ("EVM",   "Ethereum Virtual Machine"),
    ("IPFS",  "InterPlanetary File System"),
    ("MoMo",  "Mobile Money"),
    ("MVC",   "Model-View-Controller"),
    ("ORM",   "Object-Relational Mapping"),
    ("OTP",   "One-Time Password"),
    ("PDF",   "Portable Document Format"),
    ("PKI",   "Public Key Infrastructure"),
    ("PoS",   "Proof of Stake"),
    ("QR",    "Quick Response"),
    ("REST",  "Representational State Transfer"),
    ("RPC",   "Remote Procedure Call"),
    ("SDLC",  "Software Development Life Cycle"),
    ("SHA",   "Secure Hash Algorithm"),
    ("SMTP",  "Simple Mail Transfer Protocol"),
    ("SQL",   "Structured Query Language"),
    ("UML",   "Unified Modelling Language"),
    ("URL",   "Uniform Resource Locator"),
    ("XAF",   "Central African CFA Franc"),
    ("2FA",   "Two-Factor Authentication"),
]

for abbr, meaning in abbrevs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(f"{abbr}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run2 = p.add_run(f"   {meaning}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)

page_break(doc)

# ==============================================================================
# CHAPTER 1: INTRODUCTION
# ==============================================================================

add_heading1(doc, "CHAPTER 1\nINTRODUCTION")

add_heading2(doc, "1.1 Background and Context")

add_body(doc, "Education is one of the most important foundations on which individuals and societies are built. Across Africa, millions of young people pursue academic qualifications every year with the hope that those qualifications will open doors to employment, higher studies, and a better quality of life. For these qualifications to serve that purpose, they must be recognised, trusted, and verifiable by whoever receives them. Unfortunately, the systems that most African institutions use to issue and manage academic certificates have not kept pace with the scale of the challenge. The result is a growing crisis of credential fraud that affects institutions, students, employers, and governments alike.")

add_body(doc, "Credential fraud in Africa is not a minor or isolated problem. A report published by the African Union in 2019 estimated that fake certificates cost the continent billions of dollars annually through reduced productivity, poor hiring decisions, and the loss of public confidence in educational systems [1]. In Cameroon specifically, cases of forged certificates have been documented across the civil service, private sector, and even academic institutions themselves. The ease with which paper certificates can be altered, duplicated, or entirely fabricated has created a situation where an employer or university receiving a certificate has no reliable way to confirm whether it is genuine without contacting the issuing institution directly. That process is slow, unreliable, and often impossible when institutions lack proper record-keeping systems.")

add_body(doc, "At the same time, the world has been witnessing a transformation in how digital trust is established. Blockchain technology, which first gained widespread attention through the Bitcoin cryptocurrency introduced by Satoshi Nakamoto in 2008 [2], has since evolved into a general-purpose infrastructure for recording and verifying information in a way that cannot be altered after the fact. A blockchain is essentially a distributed ledger maintained by a network of computers, where each record is linked to the previous one through a cryptographic hash, making the entire chain tamper-evident. This property, known as immutability, makes blockchain an extremely attractive tool for any application where the authenticity of records is paramount.")

add_body(doc, "The idea of using blockchain for academic credentials is not new. As far back as 2015, the Massachusetts Institute of Technology began exploring digital diplomas issued on the Bitcoin blockchain through a project known as Blockcerts [3]. Since then, universities, governments, and technology companies around the world have experimented with blockchain-based credential systems. However, the overwhelming majority of these systems have been designed for and deployed in Europe, North America, and East Asia. They tend to assume reliable internet access, credit card payment systems, and English-speaking user bases. These assumptions are often not valid in the African context.")

add_body(doc, "Africa presents a unique technological environment. Mobile internet penetration has grown dramatically over the past decade, but fixed broadband remains limited in many areas. The dominant form of digital payment is not credit cards but mobile money, a system that allows transactions to be carried out using a basic mobile phone. MTN Mobile Money, known as MoMo, is one of the most widely used mobile payment platforms in Central and West Africa, with tens of millions of active users across more than 15 countries [4]. Any platform that hopes to be genuinely useful to African institutions needs to be designed with these realities in mind.")

add_body(doc, "Another important aspect of the African context is language. Many of the countries where certificate fraud is most acute are French-speaking, including Cameroon, which is officially bilingual. A certificate management platform intended for Cameroonian institutions must be accessible in French, both in its user interface and in the documents it produces. Most existing international platforms offer only English-language interfaces, which limits their adoption in Francophone Africa.")

add_body(doc, "It is against this backdrop that Certichain Africa was conceived. The platform brings together blockchain technology, distributed file storage through the InterPlanetary File System (IPFS), server-side PDF generation, mobile money payment processing, and a French-capable web interface into a single, cohesive system designed specifically for African educational institutions. The goal is not to replicate what already exists elsewhere but to build something that is genuinely adapted to the conditions and needs of the African educational landscape.")

add_heading2(doc, "1.2 Problem Statement")

add_body(doc, "The core problem that this project addresses is the absence of a reliable, accessible, and locally adapted system for issuing and verifying academic certificates in Africa, and in Cameroon in particular. To understand the depth of this problem, it is important to look at how academic certificates are currently managed and what specific weaknesses that approach creates.")

add_body(doc, "In the typical Cameroonian institution today, certificates are produced as physical documents that are printed, signed by authorised officials, and stamped with the institutional seal. These documents are then handed to the recipient, who is expected to keep them safe and present them whenever required. The institution retains a physical or digital record of who was issued a certificate, but that record is usually in the form of a spreadsheet, a printed list, or an entry in a proprietary software system that is not accessible from outside the institution.")

add_body(doc, "This approach has several serious weaknesses. First, physical certificates can be lost, damaged, or destroyed. When a student loses their original certificate, they must apply to the institution for a replacement, a process that can take weeks or months and sometimes fails entirely if the institution's records are incomplete or if key staff have left. Second, physical certificates can be forged. A sophisticated printer, the right paper, and basic image editing software are all that is needed to produce a convincing fake. Third, even genuine certificates cannot be verified remotely. An employer in Yaounde who receives a certificate issued by a university in Buea has no practical way to confirm its authenticity without making a phone call or sending a letter to the university, assuming they even know who to contact.")

add_body(doc, "The inability to verify certificates remotely is particularly damaging in the current environment, where remote work and online recruitment are becoming more common. Employers who recruit across borders or online are especially vulnerable to certificate fraud because they cannot even perform the informal checks that might be possible in a face-to-face hiring context.")

add_body(doc, "The problem is further compounded by the lack of interoperability between institutions. Each institution maintains its own records in its own format, and there is no shared infrastructure through which a verification request could be routed. Building such infrastructure at the national or continental level would require years of policy work and significant government investment. A blockchain-based approach offers an alternative path: instead of requiring institutions to connect to a central registry, it allows each institution to independently record its certificates on a shared public ledger that anyone can query.")

add_body(doc, "Existing blockchain certificate solutions do not adequately address these conditions. They are either too expensive for institutions with limited budgets, too technically complex for non-specialist staff to operate, not adapted to French-language environments, or designed around payment systems that are not available in Central Africa. As a result, most African institutions, even those that are aware of blockchain-based credential solutions, have not adopted them.")

add_body(doc, "This project therefore addresses the following specific problem: there is no web-based platform for academic certificate management that simultaneously offers blockchain-based certificate anchoring, IPFS-based storage, professional PDF generation with QR verification codes, two-factor authentication, and MTN Mobile Money payment integration, in a system designed specifically for use by educational institutions in Cameroon and the broader French-speaking African region.")

add_heading2(doc, "1.3 Objectives of the Study")

add_heading3(doc, "1.3.1 General Objective")

add_body(doc, "The general objective of this project is to design and implement a web-based platform that enables educational institutions in Cameroon and the wider African region to issue tamper-proof digital academic certificates anchored to the blockchain, manage those certificates through a secure dashboard, and make them publicly verifiable through a QR-code-based verification mechanism.")

add_heading3(doc, "1.3.2 Specific Objectives")

add_body(doc, "The following specific objectives guide the design and development of Certichain Africa:")

specifics = [
    "To design and implement a secure institution registration and authentication system that includes email-based two-factor authentication using one-time passwords.",
    "To develop a certificate creation pipeline that generates professional PDF documents for three credential types (diplomas, certifications, and badges) using the ReportLab library, with each PDF embedding a QR code for public verification.",
    "To implement a blockchain anchoring mechanism using a Solidity smart contract deployed on the Polygon network, so that each certificate is recorded on-chain in an immutable and publicly queryable manner.",
    "To integrate the Pinata IPFS service for distributed storage of certificate files, ensuring that certificate documents remain accessible independently of the platform's own servers.",
    "To build a public certificate verification system that allows any user, without logging in, to upload a certificate file or enter a blockchain hash and receive a confirmation of authenticity.",
    "To implement a tiered subscription system with payment processing through the MTN Mobile Money API, allowing institutions to pay for platform access using a payment method that is widely available in Central Africa.",
    "To design the platform with a French-capable user interface and certificate templates that reflect the linguistic and institutional context of Francophone African education.",
]

for i, obj in enumerate(specifics, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(f"{i}. {obj}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_heading2(doc, "1.4 Research Questions")

add_body(doc, "This project is guided by the following research questions, which correspond to the specific objectives stated above and which the implementation seeks to answer through its design and technical evaluation:")

questions = [
    "How can a web-based platform use blockchain technology to guarantee the authenticity and immutability of academic certificates issued by educational institutions?",
    "What combination of server-side PDF generation, QR code embedding, and public URL routing can enable instant and remote verification of academic credentials?",
    "How can the MTN Mobile Money payment API be integrated into a subscription-based platform to make the service financially accessible to institutions in Central Africa?",
    "What security mechanisms, including two-factor authentication, password hashing, rate limiting, and session management, are sufficient to protect an institution's certificate issuance authority from unauthorised access?",
    "How can a blockchain-based certificate platform be designed to remain functional even when blockchain connectivity is temporarily unavailable, through the use of a local SHA-256 fallback mechanism?",
]

for i, q in enumerate(questions, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(f"RQ{i}: {q}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_heading2(doc, "1.5 Scope and Limitations")

add_heading3(doc, "1.5.1 Scope of the Study")

add_body(doc, "Certichain Africa is scoped as an institution-facing platform. This means that the intended users of the system are educational institutions, not individual students. An institution registers on the platform, logs in through a secure authentication flow, and uses the dashboard to create and manage certificates on behalf of its students or programme graduates. The platform does not include a student-facing portal at this stage; students receive their certificates as PDF files or through a shareable verification link.")

add_body(doc, "The system covers three types of academic credentials: diplomas, which are typically awarded at the end of a degree or long-form programme; professional certifications, which attest to competence in a specific domain; and digital badges, which recognise short-form achievements or skill demonstrations. A fourth template, the Cyber Diploma, was also developed as a specialised variant for cybersecurity and technology programmes, reflecting the growing importance of this field in the region.")

add_body(doc, "The blockchain component of the system is scoped to the Polygon network, specifically because Polygon offers low transaction fees compared to the Ethereum mainnet while maintaining full compatibility with the Ethereum Virtual Machine. The IPFS component uses the Pinata cloud service rather than a self-hosted IPFS node, which reduces the operational burden for institutions while maintaining the content-addressed storage benefits of IPFS.")

add_body(doc, "The payment integration is scoped to the MTN Mobile Money Collection API, which covers Cameroon and several other Central and West African markets. The system does not currently support other payment providers such as Orange Money or card-based payments, although the architecture is designed in a way that would allow these to be added in a future iteration.")

add_heading3(doc, "1.5.2 Limitations of the Study")

add_body(doc, "Several limitations must be acknowledged. First, the platform's blockchain functionality depends on network availability and on the balance of the issuer's Polygon wallet. When the wallet runs out of MATIC tokens, the blockchain anchoring step will fail, and the system will fall back to a local SHA-256 hash instead. While this fallback ensures that the platform continues to function, it means that some certificates may not be anchored on-chain.")

add_body(doc, "Second, the IPFS storage component depends on Pinata's cloud service. If Pinata experiences downtime or if the institution's Pinata API quota is exceeded, certificate files may not be uploaded to IPFS, and the system will store them locally instead. This is a practical trade-off between simplicity of deployment and full decentralisation.")

add_body(doc, "Third, the platform does not currently include a mechanism for revoking certificates that have already been issued. Once a certificate is anchored to the blockchain, that record is permanent. Certificate revocation is an important feature for real-world deployments and is identified as a priority for future development.")

add_body(doc, "Fourth, the institution registration process does not include any form of identity verification or Know Your Customer (KYC) check. Any individual can create an institution account, which means the platform cannot guarantee that the issuing institution is legitimate. Addressing this limitation would require integration with a government identity registry or a manual vetting process.")

add_heading2(doc, "1.6 Significance of the Study")

add_body(doc, "The significance of this project can be understood from several different perspectives: technical, social, economic, and institutional.")

add_body(doc, "From a technical perspective, this project demonstrates that a blockchain-based credential management system can be built using freely available open-source tools and deployed at low cost. The combination of Flask, Web3.py, ReportLab, and Pinata that powers Certichain Africa is accessible to any developer with a basic understanding of Python web development, which means the approach can be replicated, adapted, or extended by other developers in the region without requiring specialised blockchain expertise. This is particularly important in the African context, where the pool of blockchain developers remains small.")

add_body(doc, "From a social perspective, the project directly addresses the problem of certificate fraud, which has real consequences for honest graduates who compete in job markets alongside individuals who have fraudulently obtained credentials. By making it easy for employers to verify the authenticity of a certificate in seconds using a smartphone camera and a QR code, Certichain Africa shifts the burden of proof from the honest graduate to the fraudster. A genuine certificate issued through the platform can be verified instantly; a forged one cannot.")

add_body(doc, "From an economic perspective, the project contributes to the development of the digital economy in Cameroon and the wider Central African region. The integration of MTN Mobile Money means that institutions can pay for the service using the most common digital payment method in the region, removing a key barrier to adoption. The tiered pricing model, with a free tier that allows up to five certificate issuances per month, means that small institutions and training centres can use the platform at no cost, while larger institutions pay according to their volume.")

add_body(doc, "From an institutional perspective, the project offers educational institutions a tool to protect and enhance their own reputations. An institution that issues verifiable blockchain certificates signals to the public, to employers, and to other institutions that it takes the integrity of its qualifications seriously. This can become a genuine competitive advantage in a market where institutional reputation is an important factor in student recruitment.")

add_body(doc, "Finally, from an academic perspective, this project contributes to the body of knowledge on the application of blockchain technology in the African educational context. While there is a growing literature on blockchain-based credentials globally, very little of it addresses the specific challenges of implementation in Sub-Saharan Africa. This project provides a worked example that researchers and practitioners in the region can study, critique, and build upon.")

add_heading2(doc, "1.7 Organization of the Report")

add_body(doc, "This report is organised into five main chapters, each of which builds on the previous to provide a complete account of the research and development process.")

add_body(doc, "Chapter 1, the current chapter, provides the introduction to the study. It covers the background and context, the problem statement, the objectives, the research questions, the scope and limitations, and the significance of the work.")

add_body(doc, "Chapter 2 presents the Literature Review. It begins with a conceptual review of the key technologies and ideas that underpin the project, including blockchain, IPFS, smart contracts, and digital credentials. It then discusses the theoretical framework that guided the research, followed by an empirical review of related academic work and a comparative analysis of existing platforms. The chapter concludes with an identification of the research gap that this project addresses.")

add_body(doc, "Chapter 3 describes the Methodology. It explains the research design, which combines Design Science Research with Agile development practices, and provides a detailed account of the tools and technologies selected, the system design approach, the requirements gathering process, the SDLC phases applied, and the Data Flow Diagrams that model the system's information flows.")

add_body(doc, "Chapter 4 is divided into two parts. Part 1 covers the Implementation and Results, including the full system analysis and design with diagrams, the implementation of each major feature with supporting code snippets, the testing approach, and an overview of the completed application with screenshot placeholders. Part 2 presents the Discussion, interpreting the results, comparing the system with existing works, evaluating the outcomes against the stated objectives, and reflecting on the system's strengths and limitations.")

add_body(doc, "Chapter 5 concludes the report with a summary of findings, an assessment of how well the research gap has been addressed, the conclusions drawn from the project, and recommendations for future development.")

add_body(doc, "Following the main chapters, the report includes Appendices with extended code listings and an installation guide, and a References section containing all sources cited in the report.")

page_break(doc)

# ==============================================================================
# CHAPTER 2: LITERATURE REVIEW
# ==============================================================================

add_heading1(doc, "CHAPTER 2\nLITERATURE REVIEW")

add_heading2(doc, "2.1 Introduction")

add_body(doc, "This chapter presents a systematic review of the existing body of knowledge that is relevant to the design and implementation of a blockchain-based academic certificate management system. The review is organised into five main sections. The conceptual review defines and explains the core technical and domain concepts that the project builds upon. The theoretical framework identifies the academic theories and models that guided the research design and system development approach. The empirical review examines the findings of prior studies that are directly or indirectly related to this work. The case studies section analyses existing platforms that perform similar functions. The chapter concludes with an identification of the research gap that this project is designed to fill.")

add_heading2(doc, "2.2 Conceptual Literature Review")

add_heading3(doc, "2.2.1 Digital Academic Credentials")

add_body(doc, "An academic credential is any document that attests to a person's educational achievement, qualification, or competency. In the traditional sense, credentials include degrees, diplomas, transcripts, certificates, and professional licences. The shift towards digital credentials has been accelerating over the past decade, driven by the growth of online learning, the internationalisation of labour markets, and the increasing availability of technologies that can make digital documents as trustworthy as physical ones.")

add_body(doc, "Sharma and Bhardwaj (2020) define a digital credential as an electronic representation of a qualification, skill, or achievement that is issued by a recognised authority and can be verified by a third party without the need to contact the issuing organisation directly [5]. This definition highlights two key properties that distinguish a well-designed digital credential from a simple scanned copy of a paper certificate: it must be issued by a recognised authority, and it must be independently verifiable. Both of these properties are central to the design of Certichain Africa.")

add_body(doc, "The Open Badges standard, developed by the Mozilla Foundation and later transferred to the IMS Global Learning Consortium, was one of the first widely adopted frameworks for digital credentials [6]. Open Badges are structured data packages that embed metadata about the issuer, the recipient, the criteria for the achievement, and a verification mechanism directly into the badge image or JSON structure. While Open Badges represent a significant step forward, they typically rely on centralised verification servers operated by the issuing institution, which means they are subject to the same availability and trust concerns as other centralised systems.")

add_body(doc, "The Blockcerts standard, developed at MIT, addressed this limitation by anchoring digital credentials to the Bitcoin blockchain, enabling verification without any reliance on the issuing institution's infrastructure [3]. Certichain Africa builds on this conceptual foundation but adapts it for the African institutional context, using the Polygon blockchain instead of Bitcoin, and adding features such as mobile money payment, French-language support, and server-side PDF generation that are not part of the Blockcerts standard.")

add_heading3(doc, "2.2.2 Blockchain Technology")

add_body(doc, "Blockchain is a type of distributed ledger technology in which data is stored in a chain of blocks, each containing a set of transactions or records and a cryptographic hash of the preceding block. The hash linkage between blocks means that altering any record in the chain would invalidate the hashes of all subsequent blocks, making tampering immediately detectable. Nakamoto's original Bitcoin whitepaper described this structure as a solution to the double-spending problem in digital currency [2], but the same properties that make it suitable for financial transactions also make it suitable for any application requiring an auditable, tamper-evident record.")

add_body(doc, "Buterin (2014) extended the blockchain concept by introducing Ethereum, a blockchain platform that supports general-purpose programmable logic through what he called smart contracts [7]. Smart contracts are programs stored on the blockchain that execute automatically when predetermined conditions are met. They enable the blockchain to be used not just as a ledger of transactions but as a platform for decentralised applications (DApps) that can perform complex operations without relying on a central server.")

add_body(doc, "For Certichain Africa, the relevant blockchain is Polygon, formerly known as Matic Network. Polygon is a Layer-2 scaling solution built on top of Ethereum that uses a Proof-of-Stake consensus mechanism to validate transactions much more cheaply and quickly than the Ethereum mainnet [8]. At the time of development, a certificate issuance transaction on Polygon costs a fraction of a cent in MATIC tokens, compared to several dollars on the Ethereum mainnet. This cost difference is critically important for an educational platform where the issuing institution may need to issue hundreds of certificates at a time.")

add_heading3(doc, "2.2.3 Smart Contracts")

add_body(doc, "A smart contract is a self-executing piece of code stored on a blockchain that automatically enforces the terms of an agreement when its conditions are met. Smart contracts are written in languages such as Solidity, which is the language used in this project, and are compiled into bytecode that runs on the Ethereum Virtual Machine [9].")

add_body(doc, "In the context of certificate management, a smart contract serves as the authoritative registry of issued credentials. When Certichain Africa issues a certificate, it calls the issueCertificate function on the CertiChain smart contract, passing the certificate's unique identifier (derived from the SHA-256 hash of the PDF file), the IPFS hash where the file is stored, and the recipient's name. The smart contract records this information on the blockchain, where it can be queried by anyone at any time through the verifyCertificate function. Once recorded, the entry cannot be deleted or modified, providing a permanent and public record of the issuance.")

add_body(doc, "The CertiChain.sol contract deployed in this project includes an onlyOwner modifier that restricts the issueCertificate function to the wallet address that deployed the contract. This ensures that only the platform operator can write new certificate records to the registry, preventing any unauthorised party from fabricating a blockchain record for a fake certificate.")

add_heading3(doc, "2.2.4 The InterPlanetary File System (IPFS)")

add_body(doc, "The InterPlanetary File System is a peer-to-peer distributed file storage protocol developed by Protocol Labs [10]. Unlike traditional web servers, which store files at fixed addresses (URLs) that can become inaccessible if the server goes offline, IPFS uses content-addressed storage, where each file is identified by a hash of its content, known as a Content Identifier or CID. A file's CID remains the same regardless of where the file is physically stored, and the file can be retrieved from any node in the IPFS network that has a copy of it.")

add_body(doc, "For certificate management, IPFS offers an important property: once a certificate file is uploaded to IPFS and its CID is recorded on the blockchain, the certificate can be retrieved and verified even if the platform that originally issued it ceases to exist. The blockchain record points to the IPFS CID, and the IPFS network ensures that the file remains accessible as long as at least one node has it pinned.")

add_body(doc, "In Certichain Africa, IPFS integration is provided through the Pinata cloud service, which offers an API for uploading files to IPFS and pinning them so they remain accessible indefinitely [11]. This approach avoids the need for the platform operator to run and maintain their own IPFS node, which would require significant infrastructure. The trade-off is a degree of dependence on Pinata's service availability, which is acknowledged in the limitations section of this report.")

add_heading3(doc, "2.2.5 Cryptographic Hash Functions")

add_body(doc, "A cryptographic hash function is a mathematical algorithm that takes an input of arbitrary size and produces a fixed-size output, called a hash or digest, in such a way that the same input always produces the same output, but it is computationally infeasible to reverse the process or to find two different inputs that produce the same output. These properties, known as determinism, pre-image resistance, and collision resistance respectively, make hash functions fundamental tools in both cryptography and data integrity verification [12].")

add_body(doc, "Certichain Africa uses SHA-256 hashing in two distinct ways. First, every certificate PDF file is hashed using SHA-256 when it is created, and this hash is used to generate the unique certificate identifier that is passed to the blockchain smart contract. Because SHA-256 is deterministic, any copy of the same PDF file will produce the same hash, which means verification by file upload works by simply rehashing the uploaded file and checking whether a matching record exists in the blockchain or the local database. Second, one-time passwords generated for two-factor authentication are stored as SHA-256 hashes in the server session rather than as plain text, protecting them against session hijacking attacks.")

add_body(doc, "The Solidity keccak256 function, which is the Ethereum ecosystem's native hash function, is used alongside SHA-256 in this project. Specifically, the certificate's unique identifier that is passed to the smart contract is derived using w3.solidity_keccak(['string'], [file_hash]), where file_hash is the SHA-256 digest of the certificate PDF.")

add_heading3(doc, "2.2.6 Web Application Security")

add_body(doc, "The security of a web application that handles sensitive institutional data and controls the issuance of official credentials is not an optional consideration. It is a core design requirement. Owasp (2021) identifies the top ten web application security risks, which include broken access control, cryptographic failures, injection attacks, and insecure design [13]. Each of these risks was considered in the design of Certichain Africa.")

add_body(doc, "The platform addresses broken access control through the login_required decorator applied to all protected routes, which verifies that a valid institution session exists before allowing access. Cryptographic failures are addressed through the use of Werkzeug's generate_password_hash and check_password_hash functions, which use PBKDF2-SHA256 with a random salt by default, providing strong protection against brute-force and rainbow table attacks. Injection risks are mitigated through SQLAlchemy's parameterised query interface, which prevents direct SQL injection. The OTP mechanism protects against session replay attacks by expiring codes after ten minutes and invalidating the session after five failed attempts.")

add_heading2(doc, "2.3 Theoretical Framework")

add_heading3(doc, "2.3.1 Design Science Research")

add_body(doc, "Design Science Research (DSR) is a research methodology that focuses on the creation and evaluation of novel artefacts, such as systems, models, and methods, as a means of solving real-world problems. Hevner et al. (2004) laid out a foundational framework for DSR in information systems research, arguing that a DSR project must produce an artefact that is relevant to a real business or social problem, must rigorously evaluate the artefact's performance, and must contribute to the body of knowledge in a way that extends beyond the specific problem addressed [14].")

add_body(doc, "Certichain Africa maps cleanly onto the DSR framework. The artefact is the web platform itself, including the smart contract, the PDF generation engine, the authentication system, and the mobile money payment integration. The problem addressed is the lack of a blockchain-based credential management system adapted to the African context. The evaluation takes the form of functional testing against the requirements defined in the analysis phase, and the contribution to knowledge is a worked example of how such a system can be built using accessible open-source tools.")

add_body(doc, "The six DSR activities identified by Peffers et al. (2007) also provide a useful framework for understanding the development process [15]: problem identification and motivation (recognising the certificate fraud crisis); definition of objectives for a solution (the specific objectives listed in Chapter 1); design and development (the implementation described in Chapter 4); demonstration (the running application); evaluation (the testing and validation described in Chapter 4); and communication (this report).")

add_heading3(doc, "2.3.2 The Technology Acceptance Model")

add_body(doc, "The Technology Acceptance Model (TAM), originally proposed by Davis (1989), suggests that the adoption of a new technology by a user is determined primarily by two factors: the user's perception of the technology's usefulness, and the user's perception of its ease of use [16]. When both perceived usefulness and perceived ease of use are high, a user is likely to develop a positive attitude towards the technology and adopt it.")

add_body(doc, "TAM is relevant to this project because the ultimate measure of Certichain Africa's success is not whether it works technically, but whether educational institutions actually adopt and use it. A system that is technically correct but too complex for non-technical institution administrators to use, or that does not deliver a clear benefit over the current paper-based approach, will not achieve the adoption that makes it socially valuable. This consideration influenced several design decisions, including the use of a simple, form-based certificate creation interface rather than a more technical smart contract interaction panel, and the inclusion of a free tier that allows institutions to evaluate the platform without any financial commitment.")

add_heading3(doc, "2.3.3 Diffusion of Innovations")

add_body(doc, "Rogers (2003) proposed the Diffusion of Innovations theory to explain how new ideas and technologies spread through societies [17]. According to Rogers, the rate of adoption of an innovation is influenced by five characteristics: relative advantage (how much better it is than the current approach), compatibility (how well it fits existing practices and values), complexity (how difficult it is to understand and use), trialability (whether it can be tried on a limited basis before full commitment), and observability (whether its benefits are visible to others).")

add_body(doc, "Analysing Certichain Africa through this lens reveals both strengths and challenges. The relative advantage is clear: a blockchain-anchored certificate is more verifiable, more durable, and more trustworthy than a paper certificate. Compatibility is moderate: institutions are familiar with issuing certificates but may be unfamiliar with blockchain concepts, and the platform must bridge this gap through its user interface design. Complexity is a potential barrier: even though the platform hides most of the blockchain complexity from the user, the concepts of wallets, gas fees, and IPFS are unfamiliar to most educators. Trialability is addressed by the free tier. Observability is strong: a certificate with a QR code that anyone can scan to verify it on a branded verification page is a tangible demonstration of the technology's value.")

add_heading2(doc, "2.4 Empirical Literature Review")

add_heading3(doc, "2.4.1 Blockchain in Educational Credential Systems")

add_body(doc, "A substantial body of empirical research has examined the application of blockchain technology to educational credentials. Grech and Camilleri (2017) produced a comprehensive report for the European Commission on the use of blockchain in education, identifying twenty-three potential use cases, of which digital credentials was identified as the most immediately applicable [18]. The report highlighted the advantages of blockchain credentials in terms of portability, privacy, and the elimination of institutional intermediaries in the verification process.")

add_body(doc, "Turkanovic et al. (2018) proposed EduCTX, a global higher education credit platform based on blockchain [19]. EduCTX used a permissioned blockchain to record credit transfers between universities, enabling a student who completed credits at one institution to have those credits recognised at another without the need for manual transcript verification. While EduCTX focused on credit transfer rather than final credential issuance, the underlying technical approach of using a shared ledger to enable cross-institutional recognition is directly relevant to the problem that Certichain Africa addresses.")

add_body(doc, "Lizcano et al. (2020) conducted a study of blockchain-based verification systems for academic certificates and concluded that the main barriers to adoption were the lack of standardisation in how credentials are formatted and described, the complexity of blockchain interfaces for non-technical users, and the cost of blockchain transactions in high-volume issuance scenarios [20]. These findings directly informed the design choices in Certichain Africa: the platform uses a standardised internal credential format, presents a simple form-based interface to institution users, and uses the Polygon blockchain specifically because of its low transaction costs.")

add_heading3(doc, "2.4.2 Certificate Fraud in Africa")

add_body(doc, "The empirical literature on certificate fraud in Africa is limited but consistent in its findings. Adesina (2011) documented the prevalence of fake certificates in the Nigerian labour market and found that employers had almost no effective way to verify the authenticity of academic credentials without contacting institutions directly, a process they described as time-consuming and often fruitless [21]. Adekunle et al. (2018) surveyed human resources professionals in Ghana and found that over 60 percent had encountered job applicants with suspected forged credentials at least once in the previous three years [22].")

add_body(doc, "In the Cameroonian context, reports from the Ministry of Higher Education have repeatedly identified certificate fraud as a significant challenge, particularly in the civil service recruitment process, where candidates are required to present academic certificates to qualify for positions [1]. The absence of a centralised digital registry of academic credentials means that verification relies entirely on the goodwill and administrative capacity of individual institutions, both of which are frequently insufficient.")

add_heading3(doc, "2.4.3 Mobile Money and FinTech in Africa")

add_body(doc, "MTN Mobile Money is the payment mechanism used for subscriptions in Certichain Africa. Understanding the empirical landscape of mobile money adoption in Africa is therefore important for evaluating the appropriateness of this choice. GSMA (2022) reported that Sub-Saharan Africa accounts for more than half of all mobile money accounts globally, with over 560 million registered accounts and more than 200 billion dollars in annual transaction value [4]. In Cameroon specifically, MTN MoMo is the dominant digital payment platform, with penetration rates in urban areas exceeding 70 percent among adults with mobile phones.")

add_body(doc, "Jack and Suri (2011) conducted the first major empirical study of mobile money adoption in Sub-Saharan Africa, using M-Pesa in Kenya as a case study [23]. They found that mobile money significantly increased household financial resilience and enabled economic transactions that would otherwise have been impossible due to the lack of formal banking infrastructure. Their findings support the argument that any platform targeting African institutions should prioritise mobile money integration over card-based payment systems.")

add_heading3(doc, "2.4.4 PDF Generation and Document Integrity")

add_body(doc, "The PDF format, standardised as ISO 32000, was designed by Adobe Systems to present documents in a manner that is independent of the application software, hardware, or operating system used to view them [24]. For certificate applications, PDF has several important properties: it preserves the layout and typography of the original document exactly, it can embed metadata that is not visible to the casual reader, and it can be digitally signed or hashed to detect tampering.")

add_body(doc, "ReportLab, the library used in this project for PDF generation, is a mature and widely used Python library that provides programmatic control over every aspect of a PDF document's layout, including text, graphics, images, and geometric shapes [25]. Its use in Certichain Africa allows the platform to generate professional, branded certificate documents entirely on the server side, without requiring the institution to install any software or upload pre-designed templates.")

add_heading2(doc, "2.5 Case Studies of Existing Platforms")

add_heading3(doc, "2.5.1 Blockcerts")

add_body(doc, "Blockcerts is an open standard for blockchain-based credentials developed at MIT's Media Lab and Learning Machine (now Hyland) [3]. It uses the Bitcoin blockchain to anchor credential records and provides open-source tools for issuing, displaying, and verifying Blockcerts-compliant credentials. The Blockcerts standard is well-established and has been used by MIT, the Southern New Hampshire University, and several other major institutions.")

add_body(doc, "However, Blockcerts has several characteristics that limit its applicability in the African context. First, it is a developer-focused tool rather than a ready-to-use platform, requiring significant technical expertise to deploy and operate. Second, it uses the Bitcoin blockchain, whose transaction fees, while lower than Ethereum mainnet, are still significant for high-volume issuance. Third, it does not include built-in PDF generation, payment processing, or institution management features, which means that an institution wishing to use it must build or procure those components separately.")

add_heading3(doc, "2.5.2 Accredible")

add_body(doc, "Accredible is a commercial platform for digital credential management, offering a web-based interface for issuing, managing, and sharing digital certificates and badges [26]. It is well-designed and widely used, with clients including Google, PayScale, and numerous universities. However, it is priced in US dollars, requires credit card payment, and is entirely in English. There is no blockchain anchoring of credentials; instead, Accredible relies on its own centralised servers for verification. If Accredible ceases to operate or removes a credential from its system, there is no external verification mechanism.")

add_heading3(doc, "2.5.3 Credly")

add_body(doc, "Credly, now owned by Pearson, is one of the largest digital badge platforms in the world [27]. It is based on the Open Badges standard and offers rich integration with LinkedIn and other professional networks, making it easy for credential holders to share their achievements. Like Accredible, however, Credly is centralised, priced in dollars, and not adapted to African institutional contexts. It does not offer blockchain anchoring, PDF certificate generation, or mobile money payment.")

add_heading3(doc, "2.5.4 Sony Global Education's Blockchain Credentials")

add_body(doc, "Sony Global Education partnered with IBM to develop a blockchain-based system for sharing and managing student records, using the IBM Blockchain Platform (now Hyperledger Fabric) [28]. This system is interesting because it uses a permissioned blockchain rather than a public one, which gives institutions more control over who can access the data. However, a permissioned blockchain also means that the verification is not truly open: a verifier must be granted access to the permissioned network to check a credential, which reintroduces a dependency on institutional infrastructure.")

add_heading3(doc, "2.5.5 Summary Comparison")

add_body(doc, "Table 2.1 summarises the key characteristics of the platforms reviewed above in comparison with Certichain Africa.")

tbl = doc.add_table(rows=7, cols=6)
tbl.style = 'Table Grid'
hdr = ["Feature", "Blockcerts", "Accredible", "Credly", "Sony/IBM", "Certichain Africa"]
for i, h in enumerate(hdr):
    tbl.rows[0].cells[i].text = h
    for p in tbl.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

rows_data = [
    ("Blockchain Anchoring", "Yes (Bitcoin)", "No", "No", "Yes (Permissioned)", "Yes (Polygon)"),
    ("PDF Generation", "No", "Yes", "Yes", "No", "Yes"),
    ("Mobile Money Payment", "No", "No", "No", "No", "Yes (MoMo)"),
    ("French Language Support", "Partial", "No", "No", "No", "Yes"),
    ("Free Tier Available", "Yes (self-host)", "No", "No", "No", "Yes (5 certs/month)"),
    ("African Market Focus", "No", "No", "No", "No", "Yes"),
]

for i, row_data in enumerate(rows_data, start=1):
    for j, val in enumerate(row_data):
        tbl.rows[i].cells[j].text = val
        for p in tbl.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

add_figure_placeholder(doc, "Table 2.1: Comparison of Existing Certificate Verification Platforms")

add_heading2(doc, "2.6 Research Gap")

add_body(doc, "The review of the literature and existing platforms reveals a clear and significant gap that this project addresses. The gap can be stated as follows: no existing blockchain-based academic credential platform is designed specifically for the operational, linguistic, financial, and technical constraints of educational institutions in Francophone Sub-Saharan Africa.")

add_body(doc, "This gap manifests in four specific dimensions. First, there is a financial dimension: existing platforms either require card-based payment that is not widely available in the region, or require institutions to self-host blockchain infrastructure that demands technical expertise and ongoing maintenance costs that are beyond the reach of most African institutions. Certichain Africa addresses this by offering a tiered subscription model payable through MTN Mobile Money.")

add_body(doc, "Second, there is a linguistic dimension: all major existing platforms operate primarily or exclusively in English. Certichain Africa provides a French-language interface and generates certificate documents with French text by default, reflecting the reality of Francophone African education systems.")

add_body(doc, "Third, there is a connectivity resilience dimension: existing platforms designed for markets with reliable high-speed internet do not account for situations where blockchain connectivity is temporarily unavailable. Certichain Africa addresses this with a dual-anchoring strategy that falls back to local SHA-256 fingerprinting when blockchain connectivity fails, ensuring that certificate issuance is never blocked by network issues.")

add_body(doc, "Fourth, there is a completeness dimension: no single existing platform combines blockchain anchoring, IPFS storage, server-side PDF generation with embedded QR codes, two-factor authentication, certificate management, and mobile money payment in one deployable package. Institutions that want to achieve the full set of features must either use multiple platforms, each with its own costs and integration challenges, or commission custom development. Certichain Africa provides the complete set of features in a single, cohesive platform.")

add_body(doc, "It is important to note that this project does not claim to have solved the certificate fraud problem in Africa comprehensively. The platform is a proof of concept that demonstrates the technical feasibility of a locally adapted blockchain credential system. Its wider impact depends on factors that are beyond the scope of a single Final Year Project, including institutional adoption, regulatory recognition, and national-level policy changes. Nevertheless, filling this gap with a working, well-documented system is a meaningful contribution to both the academic literature and the practical landscape of African EdTech.")

add_heading2(doc, "2.7 Conclusion")

add_body(doc, "This chapter has reviewed the conceptual foundations, theoretical frameworks, empirical evidence, and existing platforms relevant to this project. The conceptual review established an understanding of blockchain technology, smart contracts, IPFS, cryptographic hashing, and digital credentials. The theoretical framework grounded the project in Design Science Research, the Technology Acceptance Model, and Rogers' Diffusion of Innovations theory. The empirical review confirmed the reality of the certificate fraud problem in Africa, the growing adoption of mobile money, and the technical feasibility of blockchain-based credentials as demonstrated by prior research. The case study comparison showed that existing platforms all have significant limitations when considered in the African context, and the research gap section made explicit what is missing from the current landscape that this project seeks to address.")

add_body(doc, "With this understanding established, Chapter 3 turns to the methodology used to design and develop Certichain Africa, explaining the research design, the tools selected, the system design approach, and the development process.")

page_break(doc)

# ==============================================================================
# CHAPTER 3: METHODOLOGY
# ==============================================================================

add_heading1(doc, "CHAPTER 3\nMETHODOLOGY")

add_heading2(doc, "3.1 Introduction")

add_body(doc, "This chapter describes the methodological approach taken to design, develop, and evaluate Certichain Africa. Methodology in a technical project of this nature refers not only to the research philosophy that guided the work but also to the practical decisions about tools, architectures, development practices, and data collection methods. Each of these decisions has implications for the quality, reliability, and reproducibility of the results, and each is justified in the sections that follow.")

add_heading2(doc, "3.2 Research Design")

add_heading3(doc, "3.2.1 Design Science Research")

add_body(doc, "The overall research design for this project is Design Science Research (DSR). As discussed in the literature review, DSR is particularly well-suited to projects that aim to create a novel artefact (in this case, a software platform) as a means of solving a real-world problem (in this case, the lack of a blockchain-based credential system for African institutions). DSR requires that the artefact be evaluated against the requirements it was designed to meet, and that the research process produce knowledge that goes beyond the specific artefact itself.")

add_body(doc, "The DSR process for this project followed the six-stage model proposed by Peffers et al. (2007) [15]. In the problem identification and motivation stage, the certificate fraud crisis in Africa was documented and the inadequacy of existing solutions was established. In the objectives definition stage, the seven specific objectives described in Chapter 1 were formulated. In the design and development stage, the system architecture, database schema, smart contract, and user interface were designed, and the application was implemented using the technology stack described in this chapter. In the demonstration stage, the completed application was run against a set of test scenarios to show that it performs the functions it was designed to perform. In the evaluation stage, the system was assessed against the defined requirements through structured testing. In the communication stage, the design, implementation, and evaluation results are reported in this document.")

add_figure_placeholder(doc, "Figure 3.1: Design Science Research Process Applied to Certichain Africa - to be drawn")

add_heading3(doc, "3.2.2 Agile Development Approach")

add_body(doc, "Within the DSR framework, the development itself was conducted using an Agile approach. Agile is a software development philosophy that emphasises iterative development, working software over comprehensive documentation, and responsiveness to change over adherence to a fixed plan [29]. The Agile manifesto, first published in 2001, outlines four core values and twelve principles that guide Agile practice.")

add_body(doc, "For a Final Year Project developed by a single student without a formal project management structure, a lightweight version of Agile is appropriate. The development was organised into five iterative sprints, each lasting approximately two weeks. In each sprint, a set of features was designed, implemented, and tested before the next sprint began. Features that were not completed in their planned sprint were moved to the next sprint's backlog.")

add_body(doc, "The sprint plan was as follows. Sprint 1 covered the project setup, the database models, and the institution authentication system including signup, login, and OTP verification. Sprint 2 covered the certificate creation pipeline, including the PDF generation engine with all four templates and the initial blockchain integration. Sprint 3 covered the IPFS integration, the certificate management interface, and the public verification system. Sprint 4 covered the subscription pricing model, the MTN Mobile Money payment integration, and the settings management features. Sprint 5 covered the certificate templates library, error handling, security hardening, and overall testing and refinement.")

add_figure_placeholder(doc, "Figure 3.2: Agile Sprint Breakdown for Certichain Africa Development - to be drawn")

add_heading2(doc, "3.3 Tools and Technologies Used")

add_body(doc, "The selection of tools and technologies for this project was guided by three criteria: fitness for purpose (does the tool do what the project needs?), accessibility (is the tool freely available or affordable for a student project?), and community support (is there sufficient documentation and community knowledge to draw on when problems arise?). Table 3.1 summarises the full technology stack.")

tbl = doc.add_table(rows=17, cols=3)
tbl.style = 'Table Grid'
for i, h in enumerate(["Tool / Technology", "Category", "Justification"]):
    tbl.rows[0].cells[i].text = h
    for p in tbl.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

tech_rows = [
    ("Python 3.8+", "Language", "Widely taught, large standard library, strong ecosystem for web and blockchain"),
    ("Flask 2.3.3", "Web Framework", "Lightweight, minimal overhead, ideal for single-developer projects"),
    ("Flask-SQLAlchemy 3.0.5", "ORM", "Simplifies database interactions, supports multiple RDBMS backends"),
    ("SQLite / PostgreSQL", "Database", "SQLite for development (zero configuration), PostgreSQL for production"),
    ("Web3.py 6.11.0", "Blockchain Client", "Official Python library for Ethereum/Polygon interaction"),
    ("Solidity 0.8+", "Smart Contract Language", "Industry standard for EVM-compatible contracts, strong tooling"),
    ("Polygon Network", "Blockchain", "Low gas fees, EVM-compatible, active ecosystem"),
    ("Pinata Cloud API", "IPFS Storage", "Managed IPFS pinning, simple REST API, free tier available"),
    ("ReportLab 4.2.5", "PDF Generation", "Full programmatic control over PDF layout and graphics"),
    ("qrcode[pil] 7.4.2", "QR Generation", "Python library for generating QR code images embedded in PDFs"),
    ("Flask-Mail 0.9.1", "Email", "Simple SMTP integration for OTP and password reset emails"),
    ("Werkzeug 2.3.7", "Security", "PBKDF2-SHA256 password hashing, request utilities"),
    ("python-dotenv 1.0.0", "Config", "Environment variable management from .env files"),
    ("Flask-CORS 4.0.0", "API", "Cross-Origin Resource Sharing support for API endpoints"),
    ("MTN MoMo API", "Payments", "Mobile money integration for African markets"),
    ("HTML5 / CSS3 / JS", "Frontend", "Standard web technologies, no framework dependency needed"),
]

for i, (tool, cat, just) in enumerate(tech_rows, start=1):
    tbl.rows[i].cells[0].text = tool
    tbl.rows[i].cells[1].text = cat
    tbl.rows[i].cells[2].text = just
    for j in range(3):
        for p in tbl.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

add_figure_placeholder(doc, "Table 3.1: Tools and Technologies Used in Certichain Africa")

add_heading3(doc, "3.3.1 Why Flask Instead of Django")

add_body(doc, "Flask and Django are the two dominant Python web frameworks. Django is a batteries-included framework that provides a built-in ORM, admin panel, authentication system, and many other features out of the box. Flask, on the other hand, is a microframework that provides routing, request handling, and template rendering, with everything else added through extensions as needed.")

add_body(doc, "For this project, Flask was chosen because it gives the developer full control over each component of the system. Django's built-in authentication system, for example, is designed around user accounts and is not easily adapted to institution accounts with OTP two-factor authentication. Flask allows the authentication system to be designed from scratch to meet the specific requirements of the project. Flask's routing system is also simpler and more transparent than Django's URL configuration, making the codebase easier to understand and maintain for a single developer.")

add_heading3(doc, "3.3.2 Why Polygon Instead of Ethereum Mainnet")

add_body(doc, "The Ethereum mainnet is the most secure and widely used smart contract blockchain. However, at the time of development, issuing a single transaction on the Ethereum mainnet cost between two and thirty dollars in gas fees, depending on network congestion. For an educational platform where an institution might issue dozens of certificates at a time, this cost is prohibitive.")

add_body(doc, "Polygon offers full EVM compatibility, meaning that the same Solidity smart contract and the same Web3.py code that would work on Ethereum mainnet works on Polygon without modification. Transaction costs on Polygon are typically a fraction of a cent, making it economically viable for certificate issuance at any volume. Polygon uses a Proof-of-Stake consensus mechanism that makes it environmentally far less intensive than Ethereum's historical Proof-of-Work approach, which is also an ethical consideration for a project intended to serve public institutions.")

add_heading3(doc, "3.3.3 Why ReportLab for PDF Generation")

add_body(doc, "Several approaches to PDF generation in Python were considered, including WeasyPrint (which renders HTML/CSS to PDF) and PDFKit (which wraps the wkhtmltopdf command-line tool). ReportLab was chosen because it provides direct, low-level control over the canvas, allowing precise positioning of text, shapes, and images. This level of control is essential for producing the professional certificate designs required by the project, which include geometric decorations, custom fonts, brand colours, gradient effects, and precisely positioned QR codes.")

add_body(doc, "ReportLab is also a pure Python library with no external binary dependencies, which simplifies deployment on server environments where installing system libraries may not be straightforward.")

add_heading3(doc, "3.3.4 Why MTN Mobile Money")

add_body(doc, "The decision to integrate MTN Mobile Money rather than a card-based payment processor such as Stripe or PayPal reflects a deliberate design choice to make the platform accessible in markets where card payment infrastructure is limited. According to GSMA data, fewer than 20 percent of adults in Sub-Saharan Africa have a debit or credit card, while mobile money penetration is more than twice as high in many countries [4]. By accepting MoMo payments in Central African CFA Francs, the platform removes the payment barrier that would otherwise prevent most Cameroonian institutions from subscribing.")

add_heading2(doc, "3.4 System Design Approach")

add_heading3(doc, "3.4.1 MVC Architecture")

add_body(doc, "Certichain Africa is built on the Model-View-Controller (MVC) architectural pattern. In this pattern, the application is divided into three interconnected components. The Model is responsible for representing and managing the application's data. In this project, the models are defined in models.py using SQLAlchemy's declarative base and consist of three classes: Institution, Certificate, and Payment. The View is responsible for presenting data to the user. In this project, views are implemented as Jinja2 HTML templates stored in the templates directory. The Controller is responsible for handling user requests, coordinating with the model, and selecting the appropriate view. In this project, controllers are implemented as Flask route functions in app.py.")

add_body(doc, "This separation of concerns makes the codebase easier to maintain and extend. Changes to the database schema, for example, can be made in models.py without requiring changes to the templates, and changes to the visual design can be made in the templates without touching the route logic.")

add_heading3(doc, "3.4.2 RESTful API Design")

add_body(doc, "In addition to the server-side rendered page routes, Certichain Africa exposes a set of RESTful API endpoints for operations that are performed asynchronously from the frontend using JavaScript fetch calls. These endpoints are prefixed with /api/ and return JSON responses. The certificate creation API at /api/certificates/create, the certificate listing API at /api/certificates, and the payment initiation API at /api/payment/initiate are examples of this pattern. The use of JSON APIs for dynamic operations allows the frontend to update the page without a full page reload, which improves the user experience for operations like filtering certificates by type or status.")

add_heading3(doc, "3.4.3 Dual Certificate Anchoring Strategy")

add_body(doc, "One of the most important design decisions in this project was the dual anchoring strategy for certificates. The primary anchoring mechanism is the Polygon blockchain: when a certificate is created, the system attempts to call the issueCertificate function on the deployed CertiChain smart contract, passing the certificate's unique identifier, a placeholder IPFS hash, and the recipient's name. If this blockchain transaction succeeds, the certificate's blockchain_hash field in the database is set to the transaction hash prefixed with 0x, and the certificate status is set to issued.")

add_body(doc, "If the blockchain transaction fails for any reason, including insufficient wallet balance, RPC connectivity issues, or smart contract errors, the system does not fail the certificate creation. Instead, it falls back to computing a SHA-256 hash of the certificate PDF file and storing that hash prefixed with sha256: as the blockchain_hash. This fallback certificate is still verifiable through the platform's own database: any user who uploads the same PDF file will get the same SHA-256 hash, and the platform will find a matching record in the database. The fallback does not provide the same level of trust as a blockchain record, but it ensures that the certificate creation process is never blocked by blockchain infrastructure issues.")

add_heading2(doc, "3.5 Data Collection and Requirements Gathering")

add_body(doc, "The requirements for Certichain Africa were gathered through three complementary methods. First, a review of existing certificate management workflows at educational institutions was conducted by examining publicly available documentation and reports on certificate fraud in Africa, as cited in the literature review. This provided an understanding of the current state of the problem and the baseline against which the platform would need to improve.")

add_body(doc, "Second, an analysis of the capabilities and limitations of existing platforms, as described in the case studies section of the literature review, was used to identify features that should be replicated, improved upon, or avoided in the design of Certichain Africa.")

add_body(doc, "Third, a scenario-based requirements analysis was conducted, in which the typical workflows of an institution administrator using the platform were mapped out step by step. This included scenarios such as registering a new institution account, logging in for the first time with OTP verification, creating a batch of certificates for a graduating class, downloading a certificate as a PDF, and verifying a certificate uploaded by a third party. Each step in each scenario was examined to identify the data inputs, system responses, and potential failure modes that the system needed to handle.")

add_body(doc, "From this requirements gathering process, a set of functional and non-functional requirements was derived. Functional requirements describe what the system must do. Non-functional requirements describe how well the system must do it, in terms of properties like security, performance, reliability, and usability.")

add_heading2(doc, "3.6 Software Development Life Cycle Phases")

add_body(doc, "The development of Certichain Africa followed the SDLC phases adapted for an Agile project. Each phase is described below.")

add_heading3(doc, "3.6.1 Planning")

add_body(doc, "In the planning phase, the project scope, timeline, and resource requirements were defined. The choice of technology stack was made, the project repository was initialised, the Python virtual environment was set up, and the initial dependencies were installed. A rough sprint plan was drawn up, assigning features to sprints based on their dependencies and complexity.")

add_heading3(doc, "3.6.2 Analysis")

add_body(doc, "In the analysis phase, the functional and non-functional requirements were formalised. Use cases were identified and described. The database schema was designed, including the three tables for institutions, certificates, and payments, and the relationships between them. The external API integrations were analysed, and the authentication flows were mapped out.")

add_heading3(doc, "3.6.3 Design")

add_body(doc, "In the design phase, the system architecture was designed. The MVC pattern was confirmed as the structural approach. The API endpoint structure was designed, including the route paths, HTTP methods, request parameters, and response formats for each endpoint. The smart contract interface was designed, specifying the functions, parameters, and events. The PDF template designs were sketched and the brand colour palette was selected. The database models were finalised.")

add_heading3(doc, "3.6.4 Implementation")

add_body(doc, "In the implementation phase, the design was translated into working code. The implementation proceeded sprint by sprint, starting with the core authentication features and progressing through certificate management, IPFS integration, blockchain integration, payment processing, and finally testing and security hardening. Each feature was committed to the git repository after implementation, providing a version history that documents the development progress.")

add_heading3(doc, "3.6.5 Testing")

add_body(doc, "In the testing phase, the completed application was tested against the defined requirements through a combination of manual functional testing, security testing, and scenario-based end-to-end testing. Each major feature was tested through its full user journey, from input to expected output. Security-related features, including the OTP mechanism, password validation, session management, and blockchain transaction signing, were given particular attention.")

add_heading3(doc, "3.6.6 Deployment")

add_body(doc, "The deployment phase is ongoing. The application is designed to be deployable on any server that supports Python 3.8+ and can connect to the internet for blockchain RPC calls, IPFS API calls, SMTP relay, and MoMo API calls. The requirements.txt file documents all Python dependencies, and the README.md file provides step-by-step installation instructions. The application uses environment variables for all sensitive configuration, so that no secrets need to be committed to the code repository.")

add_heading2(doc, "3.7 Data Flow Diagram")

add_heading3(doc, "3.7.1 Level-0 DFD (Context Diagram)")

add_body(doc, "The Level-0 DFD, also known as the context diagram, shows Certichain Africa as a single process interacting with its external entities. The external entities are as follows:")

dfd_entities = [
    "The Institution Administrator, who interacts with the system through a web browser to register, log in, create certificates, manage certificates, and adjust settings.",
    "The Email System (SMTP server), which receives OTP codes and password reset links from the system and delivers them to the institution administrator's email address.",
    "The Polygon Blockchain Network, to which the system sends certificate issuance transactions and from which it receives transaction receipts and verification responses.",
    "The Pinata IPFS Service, to which the system uploads certificate PDF files and from which it receives IPFS Content Identifiers.",
    "The MTN MoMo API, to which the system sends payment requests and from which it receives payment status notifications.",
    "The Public User (Certificate Verifier), who interacts with the system's public verification interface without logging in to check the authenticity of a certificate.",
]
for e in dfd_entities:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(f"- {e}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_figure_placeholder(doc, "Figure 3.3: Level-0 Data Flow Diagram (Context Diagram) - to be drawn in draw.io")

add_heading3(doc, "3.7.2 Level-1 DFD")

add_body(doc, "The Level-1 DFD breaks the single Certichain Africa process into its major sub-processes. These sub-processes are as follows:")

l1_processes = [
    "Process 1 - Authentication Management: handles signup, login, OTP generation and verification, password reset, and session management. Data stores involved: Institution table in the database. External entities: Institution Administrator, Email System.",
    "Process 2 - Certificate Creation: handles the collection of certificate data from the form, PDF generation, blockchain anchoring, IPFS upload, and storage of the certificate record. Data stores: Certificate table. External entities: Polygon Blockchain, Pinata IPFS.",
    "Process 3 - Certificate Management: handles the retrieval, display, filtering, downloading, and deletion of certificates. Data stores: Certificate table, local file system.",
    "Process 4 - Certificate Verification: handles the upload of a certificate file or hash for verification, hashing the uploaded file, querying the database and blockchain, and returning the verification result. Data stores: Certificate table. External entities: Public User, Polygon Blockchain.",
    "Process 5 - Payment Processing: handles the initiation of a MoMo payment request, polling the payment status, upgrading the institution's plan on successful payment. Data stores: Payment table, Institution table. External entities: MTN MoMo API.",
    "Process 6 - Settings Management: handles updates to the institution profile, password changes, and wallet address updates. Data stores: Institution table.",
]
for proc in l1_processes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(f"- {proc}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_figure_placeholder(doc, "Figure 3.4: Level-1 Data Flow Diagram - to be drawn in draw.io")

add_heading2(doc, "3.8 Conclusion")

add_body(doc, "This chapter has described the methodological foundation of the Certichain Africa project. The research design combines Design Science Research, which provides the theoretical justification for building and evaluating a novel software artefact, with an Agile development approach, which organises the implementation into manageable, iterative sprints. The tools and technologies selected for the project were justified against the criteria of fitness for purpose, accessibility, and community support. The system design approach, based on the MVC pattern with a RESTful API layer and a dual certificate anchoring strategy, was explained. The data collection and requirements gathering methods were described, and the SDLC phases applied to the project were outlined. Finally, the Data Flow Diagrams were described in enough detail to be drawn by the student in a diagramming tool.")

add_body(doc, "With the methodology established, Chapter 4 moves into the heart of the report: the implementation of Certichain Africa, including the full system design with UML diagrams, the implementation of each major feature with supporting code snippets, the testing results, and a discussion of what the system achieves in relation to the objectives and the research gap.")

page_break(doc)

# ── Save Part 1 ──────────────────────────────────────────────────────────────
import os
os.makedirs('docs', exist_ok=True)
doc.save('docs/FYP_REPORT_PART1.docx')
print("Part 1 saved to docs/FYP_REPORT_PART1.docx")
